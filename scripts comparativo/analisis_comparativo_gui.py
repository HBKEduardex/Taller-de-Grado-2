#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
ANALIZADOR COMPARATIVO KUKA  -  KRL / MoveIt2
================================================================================

Herramienta para procesar los CSV de telemetria del KUKA
(kuka_telemetry_*.csv), calcular las metricas del estudio comparativo, generar
las graficas y guardar los resultados.

El ANALISIS (seleccionar, procesar, graficar, guardar) sigue siendo totalmente
INDEPENDIENTE de ROS/ROS2 y funciona en cualquier entorno.

La unica excepcion es la funcion OPCIONAL "Probar con RViz", que reproduce un
CSV ya procesado publicando sensor_msgs/msg/JointState en /fake_joint_states
para que RViz muestre el ensayo grabado.  Solo esa funcion necesita ROS2
(rclpy + sensor_msgs), y su import es diferido (lazy): si ROS2 no esta
disponible, la GUI arranca igual y todo lo demas sigue funcionando; unicamente
ese boton avisara de que falta el entorno ROS2.

Ejecucion:
    python3 "scripts comparativo/analisis_comparativo_gui.py"

Dependencias: pandas, numpy, matplotlib, tkinter.  scipy es OPCIONAL.
ROS2 (rclpy, sensor_msgs) es OPCIONAL: solo para "Probar con RViz".

El CSV seleccionado se abre SIEMPRE en modo lectura.  Nada se escribe hasta
pulsar el boton "Guardar resultados", y solo dentro de:
    scripts comparativo/resultados/
================================================================================
"""

import os

# DDS host <-> contenedor: RViz/ROS2 corren dentro de Docker y este Python en el
# host.  El transporte por memoria compartida de FastDDS no cruza esa frontera,
# asi que se fuerza UDPv4 y se permite el trafico fuera de localhost.  Debe
# quedar definido ANTES de cualquier import de rclpy (que es lazy, al pulsar
# "Probar con RViz").  Con setdefault, lo que el usuario exporte en la terminal
# tiene prioridad y no se pisa.
os.environ.setdefault("FASTDDS_BUILTIN_TRANSPORTS", "UDPv4")
os.environ.setdefault("ROS_LOCALHOST_ONLY", "0")

import sys
import csv
import math
import time
import subprocess
import datetime
import traceback

import numpy as np
import pandas as pd

import matplotlib
if not os.environ.get("MPLBACKEND"):
    matplotlib.use("TkAgg")
from matplotlib.figure import Figure
from matplotlib.lines import Line2D

try:
    from scipy.signal import savgol_filter as _scipy_savgol
    SCIPY_AVAILABLE = True
    import scipy as _scipy
    SCIPY_VERSION = getattr(_scipy, "__version__", "desconocida")
except Exception:
    _scipy_savgol = None
    SCIPY_AVAILABLE = False
    SCIPY_VERSION = None

# python-docx es OPCIONAL: solo lo necesita la funcion "Generar informe de varios
# CSV".  Si no esta instalado, el resto del programa funciona exactamente igual y
# ese boton avisa de que hace falta 'pip install python-docx'.
try:
    import docx as _docx
    from docx import Document as _DocxDocument
    from docx.shared import Pt as _DocxPt, Inches as _DocxInches, RGBColor as _DocxRGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _DocxAlign
    from docx.enum.section import WD_ORIENT as _DocxOrient
    DOCX_AVAILABLE = True
    DOCX_VERSION = getattr(_docx, "__version__", "desconocida")
except Exception:
    _docx = None
    _DocxDocument = None
    DOCX_AVAILABLE = False
    DOCX_VERSION = None

# Canvas Agg explicito: las figuras del informe se rasterizan desde un hilo
# secundario, y con el backend TkAgg activo savefig() crearia objetos Tk fuera
# del hilo principal.  FigureCanvasAgg evita por completo ese riesgo.
from matplotlib.backends.backend_agg import FigureCanvasAgg as _AggCanvas

from io import BytesIO
import queue as _queue
import shutil
import tempfile
import threading


# =============================================================================
# CONFIGURACION DEL ANALISIS
# (unico lugar donde se ajustan los parametros del procesamiento)
# =============================================================================

APP_TITLE = "ANALIZADOR COMPARATIVO KUKA KRL / MOVEIT2"
APP_VERSION = "1.0"

# --- Filtrado -----------------------------------------------------------------
USE_FILTER = True             # aplicar suavizado antes de derivar
SAVGOL_WINDOW = 7             # longitud de ventana (impar). ~0.9 s a 7.7 Hz
SAVGOL_POLYORDER = 3          # orden del polinomio Savitzky-Golay
FILTER_EACH_DERIVATIVE = True # re-suavizar velocidad y aceleracion antes de
                              # volver a derivar (evita jerk puramente numerico)
MOVING_AVERAGE_WINDOW = 5     # ventana del filtro alternativo si NO hay SciPy

# --- Deteccion del movimiento real -------------------------------------------
# Umbral de velocidad articular: por debajo se considera robot quieto.
MOVEMENT_VELOCITY_THRESHOLD = 0.010      # rad/s  (~0.57 deg/s)
# Recorrido articular minimo del ensayo completo para admitir que hubo movimiento.
MOVEMENT_POSITION_THRESHOLD = 0.00873    # rad    (~0.5 deg)
# Umbral adaptativo: fraccion del pico de velocidad del ensayo.
MOVEMENT_VELOCITY_PEAK_FRACTION = 0.02   # 2 % del pico
# Muestras consecutivas por encima del umbral para validar un tramo de movimiento.
MOVEMENT_MIN_CONSECUTIVE = 2
# Muestras de margen que se anaden a cada lado del tramo detectado.
MOVEMENT_EDGE_PAD_SAMPLES = 1

# --- Validacion ---------------------------------------------------------------
MIN_SAMPLES = 5               # minimo absoluto de muestras validas
MIN_DURATION_S = 1e-6         # duracion total minima admitida

# --- Graficas -----------------------------------------------------------------
PLOT_DPI = 300                # resolucion de guardado (PNG)
SCREEN_DPI = 72               # resolucion de las figuras incrustadas en la GUI
FIGSIZE_JOINTS = (13.0, 7.0)  # matriz 2x3 A1..A6
FIGSIZE_SAMPLING = (13.0, 7.0)
FIGSIZE_CARTESIAN = (13.0, 8.0)
PLOT_GRID_ALPHA = 0.35
MOTION_SPAN_COLOR = "#ffe08a"
MOTION_SPAN_ALPHA = 0.30
LINE_WIDTH = 1.4

# --- Rutas --------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
RESULTS_DIR = os.path.join(SCRIPT_DIR, "resultados")

# --- Nombres fijos ------------------------------------------------------------
AXIS_NAMES = ["A1", "A2", "A3", "A4", "A5", "A6"]
AXIS_COLUMNS = ["axis_actual.%s" % a for a in AXIS_NAMES]
TIME_COLUMN = "receive_ros_time_ns"
REQUIRED_COLUMNS = [TIME_COLUMN] + AXIS_COLUMNS

CART_POS_COLUMNS = ["position_actual.X", "position_actual.Y", "position_actual.Z"]
CART_ORI_COLUMNS = ["position_actual.A", "position_actual.B", "position_actual.C"]

FIG_FILENAMES = {
    "position":    "01_posicion_articular.png",
    "velocity":    "02_velocidad_articular.png",
    "acceleration":"03_aceleracion_articular.png",
    "jerk":        "04_jerk_articular.png",
    "continuity":  "05_continuidad_articular.png",
    "sampling":    "06_muestreo.png",
    "cart_pos":    "07_posicion_cartesiana.png",
    "cart_ori":    "08_orientacion_cartesiana.png",
}

# --- Reproduccion en RViz (FUNCION OPCIONAL: lo unico que necesita ROS2) ------
# Estos valores se copian EXACTAMENTE del nodo ya validado
#   src/kuka_telemetry_logger/kuka_telemetry_logger/rviz_mirror_node.py
# (DEFAULT_OUTPUT_TOPIC, DEFAULT_JOINT_NAMES, DEFAULT_QOS_DEPTH y su QoS
# RELIABLE / VOLATILE / KEEP_LAST).  No se inventa ninguno.
#
# NO se publica directamente en /joint_states: ese topico ya tiene publicador
# (el joint_state_publisher del visualizador, con source_list
# ["/fake_joint_states"]).  Se alimenta /fake_joint_states y se reutiliza la
# cadena existente:
#     /fake_joint_states -> joint_state_publisher -> /joint_states -> RViz
RVIZ_REPLAY_TOPIC = "/fake_joint_states"
RVIZ_REPLAY_NODE_NAME = "kuka_csv_rviz_replay"
RVIZ_REPLAY_QOS_DEPTH = 10
RVIZ_REPLAY_JOINT_NAMES = [
    "joint_a1",
    "joint_a2",
    "joint_a3",
    "joint_a4",
    "joint_a5",
    "joint_a6",
]
# Cada cuantas muestras se refresca el texto de Estado durante la reproduccion
# (no se actualiza en cada muestra para no cargar Tkinter).
RVIZ_REPLAY_PROGRESS_EVERY = 100

# Espera de descubrimiento DDS antes de arrancar la linea temporal: evita que
# se pierda la muestra 0 porque el subscriber de /fake_joint_states todavia no
# se ha emparejado con el publicador recien creado.  La espera es NO bloqueante
# (Tk.after) y NO cuenta como tiempo de reproduccion.
RVIZ_DISCOVERY_TIMEOUT_S = 2.0
RVIZ_DISCOVERY_POLL_MS = 50

ROS2_UNAVAILABLE_MESSAGE = (
    "ROS2 no está disponible en este entorno.\n"
    "Ejecute la herramienta desde una terminal donde se haya realizado\n"
    "source /opt/ros/humble/setup.bash."
)

# Metricas de la matriz de estudio que NO pueden obtenerse de la telemetria CSV.
# --- Informe grupal en Word (funcion aditiva: no afecta al analisis) ---------
# Subcarpeta exclusiva de los informes de varios CSV.  Los resultados
# individuales que ya existan en resultados/ no se tocan nunca.
GROUP_REPORTS_DIRNAME = "informes_grupales"
GROUP_REPORT_PREFIX = "informe_comparativo_grupo"
# Resolucion de las figuras incrustadas en el .docx.  Se mantiene separada de
# PLOT_DPI (300, guardado individual): a 200 dpi una figura de 13 in colocada a
# ~10.3 in de ancho conserva ~250 dpi efectivos de impresion y evita documentos
# de decenas de MB al procesar muchos CSV.  Subir a 300 si se prefiere.
REPORT_FIGURE_DPI = 200
# Documento en horizontal: las matrices 2x3 de A1-A6 y la tabla articular de 8
# columnas solo resultan legibles con el ancho de pagina apaisado.
REPORT_LANDSCAPE = True
REPORT_MARGIN_IN = 0.7
REPORT_PAGE_WIDTH_IN = 11.69          # A4 apaisado
REPORT_PAGE_HEIGHT_IN = 8.27
REPORT_BODY_FONT_PT = 10
REPORT_TABLE_FONT_PT = 9

EXTERNAL_METRICS = [
    "Error de rectitud",
    "Error angular",
    "Error de cierre",
    "Radio real de curvas",
    "Error de colocacion",
    "Repetibilidad fisica",
    "Precision fisica",
    "Tasa de exito experimental",
]


# =============================================================================
# EXCEPCIONES Y UTILIDADES
# =============================================================================

class AnalysisError(Exception):
    """Error controlado del analisis: se muestra al usuario, no cierra la app."""
    pass


def _finite(x):
    """Devuelve solo los valores finitos de un array 1D."""
    x = np.asarray(x, dtype=float).ravel()
    return x[np.isfinite(x)]


def _stats(x):
    """Estadisticos basicos de una serie: min, max, max_abs, mean, rms, std."""
    v = _finite(x)
    if v.size == 0:
        nan = float("nan")
        return {"min": nan, "max": nan, "max_abs": nan,
                "mean": nan, "rms": nan, "std": nan}
    return {
        "min": float(np.min(v)),
        "max": float(np.max(v)),
        "max_abs": float(np.max(np.abs(v))),
        "mean": float(np.mean(v)),
        "rms": float(np.sqrt(np.mean(v ** 2))),
        "std": float(np.std(v, ddof=0)),
    }


def _fmt(value, decimals=4):
    """Formatea un numero para la GUI / ficheros de texto."""
    if value is None:
        return "n/d"
    try:
        f = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not np.isfinite(f):
        return "n/d"
    return ("%." + str(decimals) + "f") % f


def _odd_window(window, n_samples, polyorder):
    """
    Ajusta una ventana de filtrado para que sea impar, <= n_samples y > polyorder.
    Devuelve None si no es posible aplicar el filtro con esos datos.
    """
    if n_samples < 3:
        return None
    w = int(window)
    if w > n_samples:
        w = n_samples
    if w % 2 == 0:
        w -= 1
    if w <= polyorder:
        return None
    if w < 3:
        return None
    return w


# =============================================================================
# 1. CARGA Y VALIDACION
# =============================================================================

def load_csv(path):
    """Lee el CSV en modo lectura. Nunca modifica el archivo original."""
    if not path:
        raise AnalysisError("No se ha seleccionado ningun archivo.")
    if not os.path.isfile(path):
        raise AnalysisError("El archivo no existe:\n%s" % path)
    if os.path.getsize(path) == 0:
        raise AnalysisError("El archivo CSV esta vacio (0 bytes):\n%s" % path)
    try:
        df = pd.read_csv(path, low_memory=False)
    except pd.errors.EmptyDataError:
        raise AnalysisError("El archivo CSV no contiene columnas ni datos:\n%s" % path)
    except pd.errors.ParserError as exc:
        raise AnalysisError("Formato CSV incorrecto.\nDetalle: %s" % exc)
    except UnicodeDecodeError:
        try:
            df = pd.read_csv(path, low_memory=False, encoding="latin-1")
        except Exception as exc:
            raise AnalysisError("No se pudo decodificar el CSV.\nDetalle: %s" % exc)
    except Exception as exc:
        raise AnalysisError("No se pudo leer el CSV.\nDetalle: %s" % exc)

    if df.shape[0] == 0:
        raise AnalysisError("El CSV no contiene filas de datos:\n%s" % path)
    return df


def validate_dataframe(df):
    """
    Comprueba que existan las columnas criticas.
    Devuelve la lista de avisos (no bloqueantes).
    """
    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        raise AnalysisError(
            "Faltan columnas criticas en el CSV.\n\n"
            "Columnas necesarias que NO se encontraron:\n  - %s\n\n"
            "El analisis articular requiere:\n  %s"
            % ("\n  - ".join(missing), ", ".join(REQUIRED_COLUMNS))
        )

    warnings = []
    for col in CART_POS_COLUMNS + CART_ORI_COLUMNS:
        if col not in df.columns:
            warnings.append("Columna cartesiana ausente: %s" % col)
    return warnings


def prepare_time(df):
    """
    Construye el eje temporal relativo en segundos a partir de receive_ros_time_ns.

    Limpia: valores no numericos, NaN/Inf, timestamps desordenados y repetidos.
    Devuelve (df_limpio, t_segundos, informe_limpieza).
    """
    report = {
        "rows_original": int(df.shape[0]),
        "dropped_non_numeric_time": 0,
        "dropped_non_numeric_axis": 0,
        "reordered": False,
        "dropped_duplicate_time": 0,
        "rows_final": 0,
    }

    work = df.copy()

    # --- tiempo -> numerico ---------------------------------------------------
    t_ns = pd.to_numeric(work[TIME_COLUMN], errors="coerce")
    bad_time = ~np.isfinite(t_ns.to_numpy(dtype=float))
    report["dropped_non_numeric_time"] = int(bad_time.sum())
    work = work.loc[~bad_time].copy()
    t_ns = t_ns.loc[~bad_time]

    if work.shape[0] == 0:
        raise AnalysisError(
            "No hay ningun timestamp valido en la columna '%s'." % TIME_COLUMN)

    # --- ejes -> numerico -----------------------------------------------------
    for col in AXIS_COLUMNS:
        work[col] = pd.to_numeric(work[col], errors="coerce")
    axis_block = work[AXIS_COLUMNS].to_numpy(dtype=float)
    bad_axis = ~np.isfinite(axis_block).all(axis=1)
    report["dropped_non_numeric_axis"] = int(bad_axis.sum())
    if bad_axis.any():
        work = work.loc[~bad_axis].copy()
        t_ns = t_ns.loc[~bad_axis]

    if work.shape[0] == 0:
        raise AnalysisError(
            "No queda ninguna muestra valida tras eliminar filas con datos "
            "articulares no numericos (NaN / Inf / texto).")

    # --- orden creciente ------------------------------------------------------
    t_values = t_ns.to_numpy(dtype=float)
    if not np.all(np.diff(t_values) >= 0):
        report["reordered"] = True
        order = np.argsort(t_values, kind="stable")
        work = work.iloc[order].copy()
        t_values = t_values[order]

    # --- timestamps repetidos (dt = 0 -> division por cero) -------------------
    keep = np.ones(t_values.shape[0], dtype=bool)
    if t_values.shape[0] > 1:
        keep[1:] = np.diff(t_values) > 0
    report["dropped_duplicate_time"] = int((~keep).sum())
    work = work.loc[keep].copy()
    t_values = t_values[keep]

    work.reset_index(drop=True, inplace=True)
    report["rows_final"] = int(work.shape[0])

    if work.shape[0] < MIN_SAMPLES:
        raise AnalysisError(
            "Muestras validas insuficientes: %d (minimo %d).\n"
            "No es posible calcular derivadas de forma fiable."
            % (work.shape[0], MIN_SAMPLES))

    t = (t_values - t_values[0]) / 1e9
    if not np.all(np.diff(t) > 0):
        raise AnalysisError("El eje temporal no es estrictamente creciente "
                            "tras la limpieza.")
    if t[-1] - t[0] <= MIN_DURATION_S:
        raise AnalysisError("La duracion total registrada es cero o despreciable.")

    return work, t, report


# =============================================================================
# 2. MUESTREO
# =============================================================================

def calculate_sampling_metrics(t):
    """
    Metricas de adquisicion.  dt en ms, duracion en s, frecuencia en Hz.
    dt_series tiene longitud N (NaN en la primera muestra) para alinear con t.
    """
    n = int(t.shape[0])
    duration = float(t[-1] - t[0])
    dt_s = np.diff(t)
    dt_ms = dt_s * 1000.0

    dt_series = np.full(n, np.nan, dtype=float)
    dt_series[1:] = dt_ms

    freq_hz = (n - 1) / duration if duration > 0 else float("nan")

    valid = dt_s[dt_s > 0]
    inst_freq = np.full(n, np.nan, dtype=float)
    if dt_s.size:
        with np.errstate(divide="ignore", invalid="ignore"):
            f_i = np.where(dt_s > 0, 1.0 / np.where(dt_s > 0, dt_s, np.nan), np.nan)
        inst_freq[1:] = f_i
    inst_freq[~np.isfinite(inst_freq)] = np.nan

    return {
        "n_samples": n,
        "duration_s": duration,
        "freq_hz": float(freq_hz),
        "dt_mean_ms": float(np.mean(dt_ms)) if dt_ms.size else float("nan"),
        "dt_median_ms": float(np.median(dt_ms)) if dt_ms.size else float("nan"),
        "dt_std_ms": float(np.std(dt_ms, ddof=0)) if dt_ms.size else float("nan"),
        "dt_min_ms": float(np.min(dt_ms)) if dt_ms.size else float("nan"),
        "dt_max_ms": float(np.max(dt_ms)) if dt_ms.size else float("nan"),
        "dt_series_ms": dt_series,
        "inst_freq_hz": inst_freq,
        "n_dt_zero_or_invalid": int(dt_s.size - valid.size),
    }


# =============================================================================
# 3. POSICIONES ARTICULARES
# =============================================================================

def prepare_joint_positions(df):
    """Extrae A1..A6.  Devuelve (grados, radianes) como arrays (N, 6)."""
    q_deg = df[AXIS_COLUMNS].to_numpy(dtype=float)
    q_rad = np.deg2rad(q_deg)
    return q_deg, q_rad


# =============================================================================
# 4. FILTRADO
# =============================================================================

def _moving_average(y, window):
    """Media movil centrada con tratamiento de bordes (fallback sin SciPy)."""
    w = int(window)
    if w < 3 or y.shape[0] < w:
        return y.copy()
    if w % 2 == 0:
        w -= 1
    half = w // 2
    padded = np.pad(y, ((half, half), (0, 0)), mode="edge")
    kernel = np.ones(w, dtype=float) / float(w)
    out = np.empty_like(y)
    for j in range(y.shape[1]):
        out[:, j] = np.convolve(padded[:, j], kernel, mode="valid")
    return out


def apply_filter(q_rad):
    """
    Suaviza las posiciones articulares ANTES de derivar.

    Nunca altera el CSV original ni el array de posiciones crudas: devuelve una
    copia suavizada.  Devuelve (q_filtrado, info) donde info documenta metodo,
    parametros y motivo si no se aplico.
    """
    n = int(q_rad.shape[0])
    info = {
        "applied": False,
        "method": "Ninguno (derivacion directa sobre datos crudos)",
        "params": "-",
        "window": None,
        "polyorder": None,
        "each_derivative": False,
        "note": "",
        "scipy_available": SCIPY_AVAILABLE,
    }

    if not USE_FILTER:
        info["note"] = "Filtrado desactivado por configuracion (USE_FILTER = False)."
        return q_rad.copy(), info

    if SCIPY_AVAILABLE:
        window = _odd_window(SAVGOL_WINDOW, n, SAVGOL_POLYORDER)
        if window is None:
            info["note"] = (
                "No se aplico Savitzky-Golay: con %d muestras no es posible una "
                "ventana impar mayor que el orden del polinomio (%d)."
                % (n, SAVGOL_POLYORDER))
            return q_rad.copy(), info
        filtered = _scipy_savgol(q_rad, window_length=window,
                                 polyorder=SAVGOL_POLYORDER, axis=0,
                                 mode="interp")
        info.update({
            "applied": True,
            "method": "Savitzky-Golay (scipy.signal.savgol_filter, mode='interp')",
            "params": "window_length=%d muestras, polyorder=%d" % (window, SAVGOL_POLYORDER),
            "window": window,
            "polyorder": SAVGOL_POLYORDER,
            "each_derivative": bool(FILTER_EACH_DERIVATIVE),
        })
        if window != SAVGOL_WINDOW:
            info["note"] = ("Ventana reducida de %d a %d por el numero de muestras "
                            "disponibles (%d)." % (SAVGOL_WINDOW, window, n))
        return filtered, info

    # --- Fallback sin SciPy ---------------------------------------------------
    window = _odd_window(MOVING_AVERAGE_WINDOW, n, 0)
    if window is None:
        info["note"] = ("SciPy no disponible y no hay muestras suficientes para la "
                        "media movil: se deriva sobre datos crudos.")
        return q_rad.copy(), info
    filtered = _moving_average(q_rad, window)
    info.update({
        "applied": True,
        "method": "Media movil centrada (alternativa: SciPy no disponible)",
        "params": "window=%d muestras" % window,
        "window": window,
        "polyorder": None,
        "each_derivative": bool(FILTER_EACH_DERIVATIVE),
        "note": "SciPy no esta instalado: se uso media movil en lugar de Savitzky-Golay.",
    })
    return filtered, info


def _smooth_like(y, info):
    """Aplica el mismo filtro configurado a una senal derivada."""
    if not info.get("applied") or not FILTER_EACH_DERIVATIVE:
        return y
    w = info.get("window")
    if not w:
        return y
    if SCIPY_AVAILABLE and info.get("polyorder") is not None:
        w2 = _odd_window(w, y.shape[0], info["polyorder"])
        if w2 is None:
            return y
        return _scipy_savgol(y, window_length=w2, polyorder=info["polyorder"],
                             axis=0, mode="interp")
    return _moving_average(y, w)


# =============================================================================
# 5. DERIVADAS
# =============================================================================

def calculate_derivatives(q_used_rad, t, filter_info):
    """
    Velocidad, aceleracion y jerk articulares con el TIEMPO REAL registrado.

    Metodo: diferencias finitas centradas de paso no uniforme
    (numpy.gradient(y, t)).  Nunca se asume dt constante.
    """
    vel = np.gradient(q_used_rad, t, axis=0)
    vel_s = _smooth_like(vel, filter_info)
    acc = np.gradient(vel_s, t, axis=0)
    acc_s = _smooth_like(acc, filter_info)
    jerk = np.gradient(acc_s, t, axis=0)
    return vel, acc, jerk


# =============================================================================
# 6. DETECCION DEL MOVIMIENTO REAL
# =============================================================================

def detect_motion(q_rad_raw, vel, t):
    """
    Determina el inicio y el fin reales del movimiento.

    Criterio:
      1) El recorrido articular total del ensayo debe superar
         MOVEMENT_POSITION_THRESHOLD (rad) en al menos una articulacion.
      2) Una muestra esta "en movimiento" si max_j |v_j| supera el umbral
         efectivo = max(MOVEMENT_VELOCITY_THRESHOLD,
                        MOVEMENT_VELOCITY_PEAK_FRACTION * pico de |v|).
      3) Solo se aceptan tramos con MOVEMENT_MIN_CONSECUTIVE muestras seguidas.
      4) Se anade un margen de MOVEMENT_EDGE_PAD_SAMPLES muestras a cada lado.
    """
    n = int(t.shape[0])
    result = {
        "detected": False,
        "i_start": 0,
        "i_end": n - 1,
        "t_start": float("nan"),
        "t_end": float("nan"),
        "duration_s": 0.0,
        "velocity_threshold": float(MOVEMENT_VELOCITY_THRESHOLD),
        "effective_velocity_threshold": float(MOVEMENT_VELOCITY_THRESHOLD),
        "position_threshold": float(MOVEMENT_POSITION_THRESHOLD),
        "max_joint_range_rad": 0.0,
        "peak_speed_rad_s": 0.0,
        "reason": "",
    }

    joint_range = np.nanmax(q_rad_raw, axis=0) - np.nanmin(q_rad_raw, axis=0)
    max_range = float(np.nanmax(joint_range)) if joint_range.size else 0.0
    result["max_joint_range_rad"] = max_range

    speed = np.nanmax(np.abs(vel), axis=1)
    peak = float(np.nanmax(speed)) if speed.size else 0.0
    result["peak_speed_rad_s"] = peak

    if max_range < MOVEMENT_POSITION_THRESHOLD:
        result["reason"] = (
            "Recorrido articular maximo %.6f rad (%.4f deg) por debajo del umbral "
            "de posicion %.6f rad." % (max_range, math.degrees(max_range),
                                       MOVEMENT_POSITION_THRESHOLD))
        return result

    thr = max(MOVEMENT_VELOCITY_THRESHOLD, MOVEMENT_VELOCITY_PEAK_FRACTION * peak)
    result["effective_velocity_threshold"] = float(thr)

    mask = np.isfinite(speed) & (speed > thr)
    if not mask.any():
        result["reason"] = ("Ninguna muestra supera el umbral de velocidad "
                            "efectivo %.6f rad/s." % thr)
        return result

    # Tramos consecutivos validos
    idx = np.flatnonzero(mask)
    runs = []
    start = idx[0]
    prev = idx[0]
    for k in idx[1:]:
        if k == prev + 1:
            prev = k
            continue
        runs.append((start, prev))
        start = k
        prev = k
    runs.append((start, prev))
    runs = [r for r in runs if (r[1] - r[0] + 1) >= MOVEMENT_MIN_CONSECUTIVE]

    if not runs:
        result["reason"] = ("Solo se detectaron picos aislados (< %d muestras "
                            "consecutivas): se interpreta como ruido."
                            % MOVEMENT_MIN_CONSECUTIVE)
        return result

    i0 = max(0, runs[0][0] - MOVEMENT_EDGE_PAD_SAMPLES)
    i1 = min(n - 1, runs[-1][1] + MOVEMENT_EDGE_PAD_SAMPLES)
    if i1 <= i0:
        result["reason"] = "Tramo de movimiento degenerado."
        return result

    result.update({
        "detected": True,
        "i_start": int(i0),
        "i_end": int(i1),
        "t_start": float(t[i0]),
        "t_end": float(t[i1]),
        "duration_s": float(t[i1] - t[i0]),
        "reason": "Movimiento detectado en %d tramo(s)." % len(runs),
    })
    return result


# =============================================================================
# 7. METRICAS ARTICULARES
# =============================================================================

def calculate_joint_metrics(q_deg, vel, acc, jerk, sl):
    """
    Metricas por articulacion dentro del intervalo de analisis `sl`.
    Posiciones en grados; velocidad rad/s, aceleracion rad/s2, jerk rad/s3.
    """
    metrics = {}
    for j, name in enumerate(AXIS_NAMES):
        pos = q_deg[sl, j]
        pos_f = _finite(pos)
        travel = float(np.sum(np.abs(np.diff(pos_f)))) if pos_f.size > 1 else 0.0
        entry = {
            "pos_inicial_deg": float(pos_f[0]) if pos_f.size else float("nan"),
            "pos_final_deg": float(pos_f[-1]) if pos_f.size else float("nan"),
            "pos_min_deg": float(np.min(pos_f)) if pos_f.size else float("nan"),
            "pos_max_deg": float(np.max(pos_f)) if pos_f.size else float("nan"),
            "recorrido_abs_deg": travel,
        }
        for key, data, unit in (("vel", vel, "rad_s"),
                                ("acc", acc, "rad_s2"),
                                ("jerk", jerk, "rad_s3")):
            st = _stats(data[sl, j])
            entry["%s_min_%s" % (key, unit)] = st["min"]
            entry["%s_max_%s" % (key, unit)] = st["max"]
            entry["%s_max_abs_%s" % (key, unit)] = st["max_abs"]
            entry["%s_mean_%s" % (key, unit)] = st["mean"]
            entry["%s_rms_%s" % (key, unit)] = st["rms"]
            entry["%s_std_%s" % (key, unit)] = st["std"]
        metrics[name] = entry
    return metrics


def calculate_continuity(q_rad_raw, sl):
    """
    Continuidad articular: delta_q[i] = |q[i] - q[i-1]| en rad, sobre las
    posiciones CRUDAS (sin filtrar) para no ocultar saltos reales entre muestras.
    Devuelve (serie_completa_Nx6_con_NaN_en_la_primera, metricas_por_eje).
    """
    n = q_rad_raw.shape[0]
    delta = np.full((n, 6), np.nan, dtype=float)
    if n > 1:
        delta[1:, :] = np.abs(np.diff(q_rad_raw, axis=0))

    metrics = {}
    for j, name in enumerate(AXIS_NAMES):
        d = _finite(delta[sl, j])
        if d.size == 0:
            nan = float("nan")
            metrics[name] = {"delta_q_max_rad": nan, "delta_q_mean_rad": nan,
                             "delta_q_rms_rad": nan, "delta_q_std_rad": nan}
            continue
        metrics[name] = {
            "delta_q_max_rad": float(np.max(d)),
            "delta_q_mean_rad": float(np.mean(d)),
            "delta_q_rms_rad": float(np.sqrt(np.mean(d ** 2))),
            "delta_q_std_rad": float(np.std(d, ddof=0)),
        }
    return delta, metrics


def calculate_cartesian_data(df):
    """Extrae X,Y,Z (mm) y A,B,C (deg) si existen.  No inventa nada."""
    data = {"position": None, "orientation": None,
            "position_missing": [], "orientation_missing": []}

    pos_missing = [c for c in CART_POS_COLUMNS if c not in df.columns]
    if not pos_missing:
        block = df[CART_POS_COLUMNS].apply(pd.to_numeric, errors="coerce").to_numpy(float)
        if np.isfinite(block).any():
            data["position"] = block
        else:
            data["position_missing"] = ["valores no numericos en X/Y/Z"]
    else:
        data["position_missing"] = pos_missing

    ori_missing = [c for c in CART_ORI_COLUMNS if c not in df.columns]
    if not ori_missing:
        block = df[CART_ORI_COLUMNS].apply(pd.to_numeric, errors="coerce").to_numpy(float)
        if np.isfinite(block).any():
            data["orientation"] = block
        else:
            data["orientation_missing"] = ["valores no numericos en A/B/C"]
    else:
        data["orientation_missing"] = ori_missing

    return data


# =============================================================================
# 8. ORQUESTADOR DEL ANALISIS  (solo memoria: no escribe nada)
# =============================================================================

class AnalysisResult(object):
    """Contenedor de todo lo calculado para un CSV.  Vive solo en memoria."""

    def __init__(self):
        self.csv_path = ""
        self.csv_name = ""
        self.processed_at = ""
        self.df = None                # dataframe limpio (copia, no el original)
        self.t = None                 # tiempo relativo [s]
        self.clean_report = {}
        self.warnings = []
        self.sampling = {}
        self.q_deg = None             # posiciones crudas [deg]
        self.q_rad = None             # posiciones crudas [rad]
        self.q_used_rad = None        # posiciones usadas para derivar [rad]
        self.filter_info = {}
        self.vel = None               # [rad/s]
        self.acc = None               # [rad/s2]
        self.jerk = None              # [rad/s3]
        self.delta_q = None           # [rad]
        self.motion = {}
        self.analysis_slice = None
        self.joint_metrics = {}
        self.continuity_metrics = {}
        self.cartesian = {}

    # -- ayudas ---------------------------------------------------------------
    @property
    def n_samples(self):
        return int(self.t.shape[0]) if self.t is not None else 0

    def metric_window_label(self):
        if self.motion.get("detected"):
            return ("Intervalo de movimiento [%.3f s , %.3f s]  (%d muestras)"
                    % (self.motion["t_start"], self.motion["t_end"],
                       self.motion["i_end"] - self.motion["i_start"] + 1))
        return ("Registro completo [%.3f s , %.3f s]  (%d muestras) - sin "
                "movimiento detectado" % (float(self.t[0]), float(self.t[-1]),
                                          self.n_samples))


def analyze_file(path):
    """Ejecuta el analisis completo de un CSV.  Devuelve un AnalysisResult."""
    raw = load_csv(path)
    warnings = validate_dataframe(raw)
    df, t, clean_report = prepare_time(raw)

    res = AnalysisResult()
    res.csv_path = os.path.abspath(path)
    res.csv_name = os.path.basename(path)
    res.processed_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    res.df = df
    res.t = t
    res.clean_report = clean_report
    res.warnings = warnings

    res.sampling = calculate_sampling_metrics(t)
    res.q_deg, res.q_rad = prepare_joint_positions(df)
    res.q_used_rad, res.filter_info = apply_filter(res.q_rad)
    res.vel, res.acc, res.jerk = calculate_derivatives(res.q_used_rad, t, res.filter_info)

    res.motion = detect_motion(res.q_rad, res.vel, t)
    if res.motion["detected"]:
        sl = slice(res.motion["i_start"], res.motion["i_end"] + 1)
    else:
        sl = slice(0, res.n_samples)
    res.analysis_slice = sl

    res.joint_metrics = calculate_joint_metrics(res.q_deg, res.vel, res.acc, res.jerk, sl)
    res.delta_q, res.continuity_metrics = calculate_continuity(res.q_rad, sl)
    res.cartesian = calculate_cartesian_data(df)

    if clean_report["dropped_non_numeric_time"]:
        warnings.append("Se descartaron %d filas con timestamp no numerico."
                        % clean_report["dropped_non_numeric_time"])
    if clean_report["dropped_non_numeric_axis"]:
        warnings.append("Se descartaron %d filas con datos articulares invalidos."
                        % clean_report["dropped_non_numeric_axis"])
    if clean_report["dropped_duplicate_time"]:
        warnings.append("Se descartaron %d filas con timestamp repetido (dt = 0)."
                        % clean_report["dropped_duplicate_time"])
    if clean_report["reordered"]:
        warnings.append("Los timestamps no estaban ordenados: se reordenaron.")
    if res.filter_info.get("note"):
        warnings.append(res.filter_info["note"])
    if not res.motion["detected"]:
        warnings.append("Sin movimiento detectado: %s" % res.motion["reason"])

    return res


# =============================================================================
# 9. GRAFICAS
# =============================================================================

def _mark_motion(ax, motion):
    """Marca inicio/fin de movimiento sin saturar la grafica."""
    if not motion.get("detected"):
        return
    ax.axvspan(motion["t_start"], motion["t_end"],
               color=MOTION_SPAN_COLOR, alpha=MOTION_SPAN_ALPHA, zorder=0)
    ax.axvline(motion["t_start"], color="#2e7d32", linestyle="--", linewidth=1.0, zorder=1)
    ax.axvline(motion["t_end"], color="#c62828", linestyle="--", linewidth=1.0, zorder=1)


def _joint_grid_figure(t, data, title, ylabel, motion, color, figsize=FIGSIZE_JOINTS,
                       mark_motion=True):
    """Figura generica 2x3 (A1..A6) frente al tiempo."""
    fig = Figure(figsize=figsize)
    axes = fig.subplots(2, 3, sharex=True)
    for j, name in enumerate(AXIS_NAMES):
        ax = axes[j // 3][j % 3]
        ax.plot(t, data[:, j], color=color, linewidth=LINE_WIDTH, label=name)
        if mark_motion:
            _mark_motion(ax, motion)
        ax.set_title(name, fontsize=11, fontweight="bold")
        ax.grid(True, alpha=PLOT_GRID_ALPHA)
        if j // 3 == 1:
            ax.set_xlabel("Tiempo [s]")
        if j % 3 == 0:
            ax.set_ylabel(ylabel)
    if mark_motion and motion.get("detected"):
        handles = [
            Line2D([], [], color="#2e7d32", linestyle="--",
                                    label="Inicio movimiento (%.2f s)" % motion["t_start"]),
            Line2D([], [], color="#c62828", linestyle="--",
                                    label="Fin movimiento (%.2f s)" % motion["t_end"]),
        ]
        fig.legend(handles=handles, loc="lower center", ncol=2, frameon=False,
                   fontsize=9, bbox_to_anchor=(0.5, 0.006))
        fig.suptitle(title, fontsize=13, fontweight="bold")
        fig.tight_layout(rect=(0, 0.045, 1, 0.965))
    else:
        fig.suptitle(title, fontsize=13, fontweight="bold")
        fig.tight_layout(rect=(0, 0, 1, 0.965))
    return fig


def create_position_figure(res):
    title = "Posicion articular A1-A6  -  %s" % res.csv_name
    return _joint_grid_figure(res.t, res.q_deg, title, "Posicion [deg]",
                              res.motion, "#1565c0")


def create_velocity_figure(res):
    title = "Velocidad articular A1-A6  -  %s" % res.csv_name
    return _joint_grid_figure(res.t, res.vel, title, "Velocidad [rad/s]",
                              res.motion, "#2e7d32")


def create_acceleration_figure(res):
    title = "Aceleracion articular A1-A6  -  %s" % res.csv_name
    return _joint_grid_figure(res.t, res.acc, title, "Aceleracion [rad/s$^2$]",
                              res.motion, "#ef6c00")


def create_jerk_figure(res):
    title = "Jerk articular A1-A6  -  %s" % res.csv_name
    return _joint_grid_figure(res.t, res.jerk, title, "Jerk [rad/s$^3$]",
                              res.motion, "#6a1b9a")


def create_continuity_figure(res):
    title = "Continuidad articular  |dq| entre muestras consecutivas  -  %s" % res.csv_name
    fig = Figure(figsize=FIGSIZE_JOINTS)
    axes = fig.subplots(2, 3, sharex=True)
    for j, name in enumerate(AXIS_NAMES):
        ax = axes[j // 3][j % 3]
        ax.plot(res.t, res.delta_q[:, j], color="#00838f", linewidth=LINE_WIDTH,
                marker="", drawstyle="steps-post")
        _mark_motion(ax, res.motion)
        mean_val = res.continuity_metrics[name]["delta_q_mean_rad"]
        if np.isfinite(mean_val):
            ax.axhline(mean_val, color="#455a64", linestyle=":", linewidth=1.0,
                       label="media = %.5f rad" % mean_val)
            ax.legend(fontsize=8, loc="upper right", framealpha=0.8)
        ax.set_title(name, fontsize=11, fontweight="bold")
        ax.grid(True, alpha=PLOT_GRID_ALPHA)
        if j // 3 == 1:
            ax.set_xlabel("Tiempo [s]")
        if j % 3 == 0:
            ax.set_ylabel("|dq| [rad]")
    fig.suptitle(title, fontsize=13, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.965))
    return fig


def create_sampling_figure(res):
    """Periodo de muestreo y frecuencia instantanea frente al tiempo."""
    s = res.sampling
    fig = Figure(figsize=FIGSIZE_SAMPLING)
    ax1, ax2 = fig.subplots(2, 1, sharex=True)

    ax1.plot(res.t, s["dt_series_ms"], color="#1565c0", linewidth=LINE_WIDTH,
             label="dt instantaneo")
    if np.isfinite(s["dt_mean_ms"]):
        ax1.axhline(s["dt_mean_ms"], color="#c62828", linestyle="--", linewidth=1.2,
                    label="dt promedio = %.2f ms" % s["dt_mean_ms"])
    if np.isfinite(s["dt_median_ms"]):
        ax1.axhline(s["dt_median_ms"], color="#2e7d32", linestyle=":", linewidth=1.2,
                    label="dt mediano = %.2f ms" % s["dt_median_ms"])
    ax1.set_ylabel("dt [ms]")
    ax1.set_title("Periodo entre muestras", fontsize=11, fontweight="bold")
    ax1.grid(True, alpha=PLOT_GRID_ALPHA)
    ax1.legend(fontsize=9, loc="best", framealpha=0.85)

    ax2.plot(res.t, s["inst_freq_hz"], color="#6a1b9a", linewidth=LINE_WIDTH,
             label="frecuencia instantanea")
    if np.isfinite(s["freq_hz"]):
        ax2.axhline(s["freq_hz"], color="#c62828", linestyle="--", linewidth=1.2,
                    label="frecuencia efectiva = %.3f Hz" % s["freq_hz"])
    ax2.set_xlabel("Tiempo [s]")
    ax2.set_ylabel("Frecuencia [Hz]")
    ax2.set_title("Frecuencia instantanea (1/dt, descartando dt = 0)",
                  fontsize=11, fontweight="bold")
    ax2.grid(True, alpha=PLOT_GRID_ALPHA)
    ax2.legend(fontsize=9, loc="best", framealpha=0.85)

    fig.suptitle("Muestreo de la telemetria  -  %s" % res.csv_name,
                 fontsize=13, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.955))
    return fig


def create_cartesian_position_figure(res):
    """Evolucion temporal de X, Y, Z [mm].  None si no existen las columnas."""
    block = res.cartesian.get("position")
    if block is None:
        return None
    fig = Figure(figsize=FIGSIZE_CARTESIAN)
    axes = fig.subplots(3, 1, sharex=True)
    labels = ["X", "Y", "Z"]
    colors = ["#1565c0", "#2e7d32", "#ef6c00"]
    for i in range(3):
        ax = axes[i]
        ax.plot(res.t, block[:, i], color=colors[i], linewidth=LINE_WIDTH,
                label=labels[i])
        _mark_motion(ax, res.motion)
        ax.set_ylabel("%s [mm]" % labels[i])
        ax.grid(True, alpha=PLOT_GRID_ALPHA)
        ax.legend(fontsize=9, loc="best", framealpha=0.85)
    axes[-1].set_xlabel("Tiempo [s]")
    fig.suptitle("Posicion cartesiana del TCP  -  %s" % res.csv_name,
                 fontsize=13, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.965))
    return fig


def create_cartesian_orientation_figure(res):
    """
    Evolucion temporal de A, B, C [deg] tal como se registran (angulos KUKA
    ZYX').  NO se deriva una velocidad angular a partir de A/B/C: esa operacion
    requiere la transformacion cinematica correcta de la representacion.
    """
    block = res.cartesian.get("orientation")
    if block is None:
        return None
    fig = Figure(figsize=FIGSIZE_CARTESIAN)
    axes = fig.subplots(3, 1, sharex=True)
    labels = ["A", "B", "C"]
    colors = ["#6a1b9a", "#00838f", "#ad1457"]
    for i in range(3):
        ax = axes[i]
        ax.plot(res.t, block[:, i], color=colors[i], linewidth=LINE_WIDTH,
                label=labels[i])
        _mark_motion(ax, res.motion)
        ax.set_ylabel("%s [deg]" % labels[i])
        ax.grid(True, alpha=PLOT_GRID_ALPHA)
        ax.legend(fontsize=9, loc="best", framealpha=0.85)
    axes[-1].set_xlabel("Tiempo [s]")
    fig.suptitle("Orientacion cartesiana registrada (A, B, C)  -  %s" % res.csv_name,
                 fontsize=13, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.965))
    return fig


def build_all_figures(res):
    """
    Construye todas las figuras.  Devuelve (dict clave->figura|None,
    lista de graficas omitidas con su motivo).
    """
    figures = {
        "position": create_position_figure(res),
        "velocity": create_velocity_figure(res),
        "acceleration": create_acceleration_figure(res),
        "jerk": create_jerk_figure(res),
        "continuity": create_continuity_figure(res),
        "sampling": create_sampling_figure(res),
        "cart_pos": create_cartesian_position_figure(res),
        "cart_ori": create_cartesian_orientation_figure(res),
    }
    skipped = []
    if figures["cart_pos"] is None:
        miss = res.cartesian.get("position_missing") or CART_POS_COLUMNS
        skipped.append((FIG_FILENAMES["cart_pos"],
                        "no disponible: %s" % ", ".join(map(str, miss))))
    if figures["cart_ori"] is None:
        miss = res.cartesian.get("orientation_missing") or CART_ORI_COLUMNS
        skipped.append((FIG_FILENAMES["cart_ori"],
                        "no disponible: %s" % ", ".join(map(str, miss))))
    return figures, skipped


# =============================================================================
# 10. CONSTRUCCION DE LOS ARCHIVOS DE RESULTADOS
# =============================================================================

def create_processed_dataframe(res):
    """
    Series temporales utilizadas en las graficas.  Archivo NUEVO: el CSV
    original nunca se modifica.
    """
    data = {
        "time_s": res.t,
        "dt_ms": res.sampling["dt_series_ms"],
        "instant_freq_hz": res.sampling["inst_freq_hz"],
    }
    in_motion = np.zeros(res.n_samples, dtype=int)
    if res.motion.get("detected"):
        in_motion[res.motion["i_start"]:res.motion["i_end"] + 1] = 1
    data["in_motion_window"] = in_motion

    for j, name in enumerate(AXIS_NAMES):
        data["%s_position_deg" % name] = res.q_deg[:, j]
    for j, name in enumerate(AXIS_NAMES):
        data["%s_position_rad" % name] = res.q_rad[:, j]
    if res.filter_info.get("applied"):
        for j, name in enumerate(AXIS_NAMES):
            data["%s_position_filtered_rad" % name] = res.q_used_rad[:, j]
    for j, name in enumerate(AXIS_NAMES):
        data["%s_velocity_rad_s" % name] = res.vel[:, j]
    for j, name in enumerate(AXIS_NAMES):
        data["%s_acceleration_rad_s2" % name] = res.acc[:, j]
    for j, name in enumerate(AXIS_NAMES):
        data["%s_jerk_rad_s3" % name] = res.jerk[:, j]
    for j, name in enumerate(AXIS_NAMES):
        data["%s_delta_q_rad" % name] = res.delta_q[:, j]

    block = res.cartesian.get("position")
    if block is not None:
        for i, label in enumerate(["X", "Y", "Z"]):
            data["%s_mm" % label] = block[:, i]
    block = res.cartesian.get("orientation")
    if block is not None:
        for i, label in enumerate(["A", "B", "C"]):
            data["%s_deg" % label] = block[:, i]

    return pd.DataFrame(data)


def create_metrics_dataframe(res):
    """
    resumen_metricas.csv: una fila por articulacion.  Las columnas generales del
    ensayo se repiten en cada fila para que el archivo sea rectangular y se abra
    sin problemas en Excel, Python o MATLAB.
    """
    s = res.sampling
    m = res.motion
    general = {
        "archivo_original": res.csv_name,
        "ruta_original": res.csv_path,
        "fecha_procesamiento": res.processed_at,
        "numero_muestras": s["n_samples"],
        "duracion_total_s": s["duration_s"],
        "duracion_movimiento_s": m["duration_s"] if m["detected"] else 0.0,
        "inicio_movimiento_s": m["t_start"] if m["detected"] else float("nan"),
        "fin_movimiento_s": m["t_end"] if m["detected"] else float("nan"),
        "frecuencia_efectiva_hz": s["freq_hz"],
        "dt_promedio_ms": s["dt_mean_ms"],
        "dt_mediano_ms": s["dt_median_ms"],
        "dt_std_ms": s["dt_std_ms"],
        "dt_min_ms": s["dt_min_ms"],
        "dt_max_ms": s["dt_max_ms"],
        "movimiento_detectado": "Si" if m["detected"] else "No",
        "intervalo_metricas": ("movimiento" if m["detected"] else "registro_completo"),
        "umbral_velocidad_rad_s": m["effective_velocity_threshold"],
        "umbral_posicion_rad": m["position_threshold"],
        "filtrado_aplicado": "Si" if res.filter_info.get("applied") else "No",
        "metodo_filtrado": res.filter_info.get("method", "-"),
        "parametros_filtrado": res.filter_info.get("params", "-"),
        "metodo_derivacion": "numpy.gradient sobre tiempo real no uniforme",
    }

    rows = []
    for name in AXIS_NAMES:
        row = dict(general)
        row["axis"] = name
        row.update(res.joint_metrics[name])
        row.update(res.continuity_metrics[name])
        rows.append(row)

    columns = (["archivo_original", "ruta_original", "fecha_procesamiento",
                "numero_muestras", "duracion_total_s", "duracion_movimiento_s",
                "inicio_movimiento_s", "fin_movimiento_s", "frecuencia_efectiva_hz",
                "dt_promedio_ms", "dt_mediano_ms", "dt_std_ms", "dt_min_ms",
                "dt_max_ms", "movimiento_detectado", "intervalo_metricas",
                "umbral_velocidad_rad_s", "umbral_posicion_rad",
                "filtrado_aplicado", "metodo_filtrado", "parametros_filtrado",
                "metodo_derivacion", "axis",
                "pos_inicial_deg", "pos_final_deg", "pos_min_deg", "pos_max_deg",
                "recorrido_abs_deg"]
               + ["vel_%s_rad_s" % k for k in ("min", "max", "max_abs", "mean", "rms", "std")]
               + ["acc_%s_rad_s2" % k for k in ("min", "max", "max_abs", "mean", "rms", "std")]
               + ["jerk_%s_rad_s3" % k for k in ("min", "max", "max_abs", "mean", "rms", "std")]
               + ["delta_q_max_rad", "delta_q_mean_rad", "delta_q_rms_rad",
                  "delta_q_std_rad"])
    return pd.DataFrame(rows)[columns]


def create_summary_txt(res, skipped_figures=None):
    """Texto legible por una persona.  Se usa en la GUI y en resumen.txt."""
    s = res.sampling
    m = res.motion
    fi = res.filter_info
    L = []
    add = L.append

    add("=" * 72)
    add("ESTUDIO COMPARATIVO KUKA  -  ANALISIS DE TELEMETRIA")
    add("=" * 72)
    add("")
    add("Archivo analizado:")
    add("  %s" % res.csv_name)
    add("  %s" % res.csv_path)
    add("Fecha y hora del procesamiento: %s" % res.processed_at)
    add("Herramienta: analisis_comparativo_gui.py v%s" % APP_VERSION)
    add("")
    add("DATOS DE ADQUISICION")
    add("-" * 72)
    add("Numero de muestras validas:   %d" % s["n_samples"])
    add("Filas del CSV original:       %d" % res.clean_report["rows_original"])
    add("Duracion total registrada:    %s s" % _fmt(s["duration_s"], 3))
    add("Frecuencia efectiva:          %s Hz" % _fmt(s["freq_hz"], 3))
    add("Periodo promedio:             %s ms" % _fmt(s["dt_mean_ms"], 3))
    add("Periodo mediano:              %s ms" % _fmt(s["dt_median_ms"], 3))
    add("Desviacion del periodo:       %s ms" % _fmt(s["dt_std_ms"], 3))
    add("Periodo minimo:               %s ms" % _fmt(s["dt_min_ms"], 3))
    add("Periodo maximo:               %s ms" % _fmt(s["dt_max_ms"], 3))
    add("")
    add("LIMPIEZA DE DATOS")
    add("-" * 72)
    add("Filas con timestamp no numerico descartadas: %d"
        % res.clean_report["dropped_non_numeric_time"])
    add("Filas con datos articulares invalidos descartadas: %d"
        % res.clean_report["dropped_non_numeric_axis"])
    add("Filas con timestamp repetido descartadas:    %d"
        % res.clean_report["dropped_duplicate_time"])
    add("Timestamps reordenados:                      %s"
        % ("Si" if res.clean_report["reordered"] else "No"))
    add("")
    add("MOVIMIENTO")
    add("-" * 72)
    add("Movimiento detectado:     %s" % ("Si" if m["detected"] else "No"))
    if m["detected"]:
        add("Inicio movimiento:        %s s" % _fmt(m["t_start"], 3))
        add("Fin movimiento:           %s s" % _fmt(m["t_end"], 3))
        add("Duracion movimiento:      %s s" % _fmt(m["duration_s"], 3))
    else:
        add("Inicio movimiento:        n/d")
        add("Fin movimiento:           n/d")
        add("Duracion movimiento:      0.000 s")
        add("Motivo: %s" % m["reason"])
        add("NOTA: no se atribuye velocidad, aceleracion, jerk ni tiempo de")
        add("      ejecucion significativos a este ensayo.  Las cifras de las")
        add("      tablas corresponden al ruido residual del registro completo.")
    add("Recorrido articular maximo del ensayo: %s rad (%s deg)"
        % (_fmt(m["max_joint_range_rad"], 6),
           _fmt(math.degrees(m["max_joint_range_rad"]), 4)))
    add("Pico de velocidad articular:           %s rad/s" % _fmt(m["peak_speed_rad_s"], 6))
    add("")
    add("PROCESAMIENTO  (trazabilidad)")
    add("-" * 72)
    add("Eje temporal:             columna '%s', tiempo relativo en segundos" % TIME_COLUMN)
    add("Metodo de derivacion:     numpy.gradient(y, t) - diferencias finitas")
    add("                          centradas con paso NO uniforme (tiempo real).")
    add("                          velocidad = dq/dt, aceleracion = dv/dt, jerk = da/dt")
    add("Filtrado aplicado:        %s" % ("Si" if fi.get("applied") else "No"))
    add("Metodo de filtrado:       %s" % fi.get("method", "-"))
    add("Parametros del filtrado:  %s" % fi.get("params", "-"))
    add("Filtro en cada derivada:  %s" % ("Si" if fi.get("each_derivative") else "No"))
    add("SciPy disponible:         %s%s"
        % ("Si" if SCIPY_AVAILABLE else "No",
           (" (v%s)" % SCIPY_VERSION) if SCIPY_AVAILABLE else ""))
    if fi.get("note"):
        add("Observacion del filtrado: %s" % fi["note"])
    add("Umbral de velocidad para detectar movimiento: %s rad/s"
        % _fmt(m["effective_velocity_threshold"], 6))
    add("  (configurado: %s rad/s ; adaptativo: %.1f %% del pico)"
        % (_fmt(MOVEMENT_VELOCITY_THRESHOLD, 6),
           100.0 * MOVEMENT_VELOCITY_PEAK_FRACTION))
    add("Umbral de recorrido para admitir movimiento:  %s rad"
        % _fmt(MOVEMENT_POSITION_THRESHOLD, 6))
    add("Muestras consecutivas minimas:                %d" % MOVEMENT_MIN_CONSECUTIVE)
    add("Margen anadido al tramo detectado:            %d muestras"
        % MOVEMENT_EDGE_PAD_SAMPLES)
    add("Continuidad |dq|: calculada sobre las posiciones CRUDAS (sin filtrar).")
    add("")
    add("INTERVALO UTILIZADO PARA LAS METRICAS")
    add("-" * 72)
    add(res.metric_window_label())
    add("")
    add("POSICION ARTICULAR  [deg]")
    add("-" * 72)
    add("%-4s %12s %12s %12s %12s %14s"
        % ("Eje", "inicial", "final", "minima", "maxima", "recorrido abs"))
    for name in AXIS_NAMES:
        e = res.joint_metrics[name]
        add("%-4s %12s %12s %12s %12s %14s"
            % (name, _fmt(e["pos_inicial_deg"], 4), _fmt(e["pos_final_deg"], 4),
               _fmt(e["pos_min_deg"], 4), _fmt(e["pos_max_deg"], 4),
               _fmt(e["recorrido_abs_deg"], 4)))
    add("")
    add("METRICAS PRINCIPALES POR ARTICULACION")
    add("-" * 72)
    for name in AXIS_NAMES:
        e = res.joint_metrics[name]
        c = res.continuity_metrics[name]
        add("%s:" % name)
        add("  Velocidad    [rad/s]   max_abs=%s  rms=%s  mean=%s  std=%s  min=%s  max=%s"
            % (_fmt(e["vel_max_abs_rad_s"], 6), _fmt(e["vel_rms_rad_s"], 6),
               _fmt(e["vel_mean_rad_s"], 6), _fmt(e["vel_std_rad_s"], 6),
               _fmt(e["vel_min_rad_s"], 6), _fmt(e["vel_max_rad_s"], 6)))
        add("  Aceleracion  [rad/s2]  max_abs=%s  rms=%s  mean=%s  std=%s  min=%s  max=%s"
            % (_fmt(e["acc_max_abs_rad_s2"], 6), _fmt(e["acc_rms_rad_s2"], 6),
               _fmt(e["acc_mean_rad_s2"], 6), _fmt(e["acc_std_rad_s2"], 6),
               _fmt(e["acc_min_rad_s2"], 6), _fmt(e["acc_max_rad_s2"], 6)))
        add("  Jerk         [rad/s3]  max_abs=%s  rms=%s  mean=%s  std=%s  min=%s  max=%s"
            % (_fmt(e["jerk_max_abs_rad_s3"], 6), _fmt(e["jerk_rms_rad_s3"], 6),
               _fmt(e["jerk_mean_rad_s3"], 6), _fmt(e["jerk_std_rad_s3"], 6),
               _fmt(e["jerk_min_rad_s3"], 6), _fmt(e["jerk_max_rad_s3"], 6)))
        add("  Continuidad  [rad]     dq_max=%s  dq_rms=%s  dq_mean=%s  dq_std=%s"
            % (_fmt(c["delta_q_max_rad"], 6), _fmt(c["delta_q_rms_rad"], 6),
               _fmt(c["delta_q_mean_rad"], 6), _fmt(c["delta_q_std_rad"], 6)))
        add("")

    add("DATOS CARTESIANOS")
    add("-" * 72)
    if res.cartesian.get("position") is not None:
        block = res.cartesian["position"]
        for i, label in enumerate(["X", "Y", "Z"]):
            col = _finite(block[:, i])
            if col.size:
                add("  %s [mm]: inicial=%s  final=%s  min=%s  max=%s"
                    % (label, _fmt(col[0], 3), _fmt(col[-1], 3),
                       _fmt(np.min(col), 3), _fmt(np.max(col), 3)))
    else:
        add("  Posicion cartesiana no disponible: %s"
            % ", ".join(map(str, res.cartesian.get("position_missing") or [])))
    if res.cartesian.get("orientation") is not None:
        block = res.cartesian["orientation"]
        for i, label in enumerate(["A", "B", "C"]):
            col = _finite(block[:, i])
            if col.size:
                add("  %s [deg]: inicial=%s  final=%s  min=%s  max=%s"
                    % (label, _fmt(col[0], 3), _fmt(col[-1], 3),
                       _fmt(np.min(col), 3), _fmt(np.max(col), 3)))
        add("  (No se deriva velocidad angular del efector a partir de A/B/C:")
        add("   requiere la transformacion cinematica correcta de esa")
        add("   representacion de orientacion.)")
    else:
        add("  Orientacion cartesiana no disponible: %s"
            % ", ".join(map(str, res.cartesian.get("orientation_missing") or [])))
    add("")

    if skipped_figures:
        add("GRAFICAS NO GENERADAS")
        add("-" * 72)
        for fname, reason in skipped_figures:
            add("  %s : %s" % (fname, reason))
        add("")

    if res.warnings:
        add("AVISOS")
        add("-" * 72)
        for w in res.warnings:
            add("  - %s" % w)
        add("")

    add("LIMITACIONES")
    add("-" * 72)
    add("Indice completo de singularidad: NO calculado.")
    add("  Se requiere el modelo cinematico / Jacobiano del KUKA KR6 R900")
    add("  (parametros DH o URDF, dimensiones geometricas).  Con las columnas")
    add("  disponibles en el CSV no existe informacion suficiente y no se")
    add("  implementa ninguna formula arbitraria.")
    add("")
    add("Metricas que requieren medicion experimental externa (no se calculan")
    add("a partir de la telemetria y no se inventan):")
    for item in EXTERNAL_METRICS:
        add("  - %s : requiere medicion experimental externa." % item)
    add("")
    add("=" * 72)
    add("FIN DEL INFORME")
    add("=" * 72)
    return "\n".join(L)


# =============================================================================
# 11. GUARDADO
# =============================================================================

def _unique_run_folder(base_dir, stem):
    """
    Devuelve una carpeta nueva para el ensayo.  Nunca sobrescribe resultados
    anteriores: anade _01, _02, ... y como ultimo recurso fecha y hora.
    """
    candidate = os.path.join(base_dir, stem)
    if not os.path.exists(candidate):
        return candidate
    for i in range(1, 100):
        candidate = os.path.join(base_dir, "%s_%02d" % (stem, i))
        if not os.path.exists(candidate):
            return candidate
    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(base_dir, "%s_%s" % (stem, stamp))


def save_results(res, figures, skipped_figures):
    """
    Escribe todos los resultados en scripts comparativo/resultados/<ensayo>/.
    Solo se llama desde el boton "Guardar resultados".
    Devuelve (carpeta, lista_de_archivos_escritos).
    """
    os.makedirs(RESULTS_DIR, exist_ok=True)
    stem = os.path.splitext(res.csv_name)[0]
    folder = _unique_run_folder(RESULTS_DIR, stem)
    os.makedirs(folder, exist_ok=False)

    written = []

    metrics_df = create_metrics_dataframe(res)
    path = os.path.join(folder, "resumen_metricas.csv")
    metrics_df.to_csv(path, index=False, quoting=csv.QUOTE_MINIMAL)
    written.append(os.path.basename(path))

    processed_df = create_processed_dataframe(res)
    path = os.path.join(folder, "datos_procesados.csv")
    processed_df.to_csv(path, index=False)
    written.append(os.path.basename(path))

    for key, filename in FIG_FILENAMES.items():
        fig = figures.get(key)
        if fig is None:
            continue
        target = os.path.join(folder, filename)
        if key == "sampling":
            size = FIGSIZE_SAMPLING
        elif key in ("cart_pos", "cart_ori"):
            size = FIGSIZE_CARTESIAN
        else:
            size = FIGSIZE_JOINTS
        old_size = fig.get_size_inches().copy()
        try:
            fig.set_size_inches(size[0], size[1])
            fig.savefig(target, dpi=PLOT_DPI, bbox_inches="tight")
        finally:
            fig.set_size_inches(old_size[0], old_size[1])
        written.append(filename)

    text = create_summary_txt(res, skipped_figures)
    path = os.path.join(folder, "resumen.txt")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(text)
        fh.write("\n")
    written.append(os.path.basename(path))

    return folder, written


def open_results_folder(folder):
    """Abre la carpeta en el explorador del sistema, sin usar shell."""
    if not folder or not os.path.isdir(folder):
        raise AnalysisError("La carpeta de resultados no existe:\n%s" % folder)
    try:
        if sys.platform.startswith("linux"):
            subprocess.Popen(["xdg-open", folder],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", folder],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif os.name == "nt":
            os.startfile(folder)  # noqa: F821  (solo existe en Windows)
        else:
            raise AnalysisError("Plataforma no soportada para abrir carpetas: %s"
                                % sys.platform)
    except FileNotFoundError:
        raise AnalysisError("No se encontro el gestor de archivos del sistema.\n"
                            "Carpeta:\n%s" % folder)
    except OSError as exc:
        raise AnalysisError("No se pudo abrir la carpeta.\nDetalle: %s" % exc)


# =============================================================================
# 12-bis. INFORME GRUPAL EN WORD  (funcion ADITIVA)
#
# Procesa varios CSV en una sola operacion y construye UN unico .docx.
#
# Reutiliza integramente el motor de analisis del flujo individual:
#   analyze_file()  ->  build_all_figures()
# No hay un segundo algoritmo: las metricas y las figuras del informe son
# exactamente las mismas que muestra y guarda el analisis de un solo CSV.
# =============================================================================

# Metricas articulares que se agregan entre ensayos (clave, etiqueta, unidad, decimales)
GROUP_AXIS_METRICS = [
    ("vel_max_abs_rad_s", "Velocidad maxima absoluta", "rad/s", 5),
    ("vel_rms_rad_s", "Velocidad RMS", "rad/s", 5),
    ("acc_max_abs_rad_s2", "Aceleracion maxima absoluta", "rad/s2", 5),
    ("acc_rms_rad_s2", "Aceleracion RMS", "rad/s2", 5),
    ("jerk_max_abs_rad_s3", "Jerk maximo absoluto", "rad/s3", 5),
    ("jerk_rms_rad_s3", "Jerk RMS", "rad/s3", 5),
    ("delta_q_max_rad", "Delta q maximo", "rad", 6),
]

# Orden de las figuras dentro de la seccion de cada ensayo
GROUP_FIGURE_ORDER = [
    ("position", "Posicion articular", "Posicion articular"),
    ("velocity", "Velocidad articular", "Velocidad articular"),
    ("acceleration", "Aceleracion articular", "Aceleracion articular"),
    ("jerk", "Jerk articular", "Jerk articular"),
    ("continuity", "Continuidad articular", "Continuidad articular"),
    ("sampling", "Muestreo", "Comportamiento del periodo de muestreo"),
    ("cart_pos", "Posicion cartesiana", "Posicion cartesiana del TCP"),
    ("cart_ori", "Orientacion cartesiana", "Orientacion cartesiana registrada"),
]


def docx_requirement_message():
    """Mensaje unico para cuando falta python-docx."""
    return ("Esta funcion necesita la libreria 'python-docx', que no esta "
            "instalada.\n\nInstalela con:\n\n    pip install python-docx\n\n"
            "El resto del programa funciona con normalidad sin ella.")


def group_reports_dir():
    """scripts comparativo/resultados/informes_grupales/ (se crea al vuelo)."""
    return os.path.join(RESULTS_DIR, GROUP_REPORTS_DIRNAME)


def group_report_path(when=None):
    """informe_comparativo_grupo_YYYYMMDD_HHMMSS.docx (nunca sobrescribe)."""
    when = when or datetime.datetime.now()
    name = "%s_%s.docx" % (GROUP_REPORT_PREFIX, when.strftime("%Y%m%d_%H%M%S"))
    return os.path.join(group_reports_dir(), name)


def report_usable_width_in():
    """Ancho util de la pagina del informe, en pulgadas."""
    width = REPORT_PAGE_WIDTH_IN if REPORT_LANDSCAPE else REPORT_PAGE_HEIGHT_IN
    return width - 2 * REPORT_MARGIN_IN


# -----------------------------------------------------------------------------
# Figuras -> PNG temporal
# -----------------------------------------------------------------------------

def _release_figure(fig):
    """Libera una figura de Matplotlib (equivalente a plt.close para Figure)."""
    if fig is None:
        return
    try:
        fig.clear()
    except Exception:
        pass
    try:
        fig.canvas = None
    except Exception:
        pass


def render_figure_to_png(fig, target_path):
    """
    Rasteriza una figura a PNG usando explicitamente el canvas Agg.

    Se fuerza Agg porque esta funcion se ejecuta en un hilo secundario y el
    backend activo de la aplicacion es TkAgg (no seguro fuera del hilo principal).
    """
    _AggCanvas(fig)
    fig.savefig(target_path, format="png", dpi=REPORT_FIGURE_DPI,
                bbox_inches="tight")
    return target_path


def render_figures_to_tempdir(figures, tmp_dir, test_index):
    """
    Guarda las figuras del ensayo en PNG temporales y devuelve {clave: ruta}.
    Los PNG viven en una carpeta temporal que se borra al terminar el informe:
    no se crean imagenes permanentes dentro de resultados/.
    """
    paths = {}
    for key, _sub, _cap in GROUP_FIGURE_ORDER:
        fig = figures.get(key)
        if fig is None:
            continue
        target = os.path.join(tmp_dir, "e%02d_%s.png" % (test_index, key))
        try:
            render_figure_to_png(fig, target)
            paths[key] = target
        except Exception:
            # una figura que falle no debe tumbar el informe completo
            continue
    return paths


def add_matplotlib_figure_to_docx(document, image_path, width_in=None):
    """Inserta un PNG centrado, ajustado al ancho util manteniendo proporcion."""
    if width_in is None:
        width_in = report_usable_width_in()
    par = document.add_paragraph()
    par.alignment = _DocxAlign.CENTER
    run = par.add_run()
    run.add_picture(image_path, width=_DocxInches(width_in))
    return par


def add_figure_caption(document, text):
    """Pie de figura centrado y en cursiva."""
    par = document.add_paragraph()
    par.alignment = _DocxAlign.CENTER
    run = par.add_run(text)
    run.italic = True
    run.font.size = _DocxPt(REPORT_TABLE_FONT_PT)
    return par


# -----------------------------------------------------------------------------
# Utilidades de documento
# -----------------------------------------------------------------------------

def _new_report_document():
    """Documento base: A4 apaisado, margenes ajustados y fuente legible."""
    document = _DocxDocument()
    section = document.sections[0]
    if REPORT_LANDSCAPE:
        section.orientation = _DocxOrient.LANDSCAPE
        section.page_width = _DocxInches(REPORT_PAGE_WIDTH_IN)
        section.page_height = _DocxInches(REPORT_PAGE_HEIGHT_IN)
    else:
        section.page_width = _DocxInches(REPORT_PAGE_HEIGHT_IN)
        section.page_height = _DocxInches(REPORT_PAGE_WIDTH_IN)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, _DocxInches(REPORT_MARGIN_IN))
    try:
        document.styles["Normal"].font.size = _DocxPt(REPORT_BODY_FONT_PT)
    except Exception:
        pass
    return document


def _add_paragraph(document, text="", bold=False, size_pt=None, align=None):
    par = document.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    if size_pt is not None:
        run.font.size = _DocxPt(size_pt)
    return par


def _set_fixed_table_widths(table, widths_in):
    """
    Fija los anchos de columna de verdad.

    Word y LibreOffice reajustan las columnas por su cuenta salvo que la tabla
    declare 'layout fijo'; sin esto, los nombres largos de archivo parten en
    varias lineas y la tabla queda ilegible.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    try:
        table.autofit = False
    except Exception:
        pass
    try:
        tbl_pr = table._tbl.tblPr
        for existing in tbl_pr.findall(qn("w:tblLayout")):
            tbl_pr.remove(existing)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
    except Exception:
        pass

    for i, width in enumerate(widths_in):
        if i >= len(table.columns):
            break
        try:
            table.columns[i].width = _DocxInches(width)
        except Exception:
            pass
    for row in table.rows:
        for i, width in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = _DocxInches(width)


def _add_table(document, headers, rows, widths_in=None,
               font_pt=REPORT_TABLE_FONT_PT):
    """
    Tabla con cabecera en negrita.  Los valores van como TEXTO real editable,
    nunca como imagen.
    """
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        header_cells[i].text = ""
        run = header_cells[i].paragraphs[0].add_run(str(text))
        run.bold = True
        run.font.size = _DocxPt(font_pt)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = _DocxPt(font_pt)
    if widths_in:
        _set_fixed_table_widths(table, widths_in)
    return table


def _add_page_break(document):
    document.add_page_break()


def _mean_std_min_max(values):
    """Estadisticos entre ensayos, ignorando valores no finitos."""
    v = _finite(values)
    if v.size == 0:
        nan = float("nan")
        return {"n": 0, "mean": nan, "std": nan, "min": nan, "max": nan}
    return {"n": int(v.size), "mean": float(np.mean(v)),
            "std": float(np.std(v, ddof=0)), "min": float(np.min(v)),
            "max": float(np.max(v))}


# -----------------------------------------------------------------------------
# Resumen compacto de un ensayo (lo unico que se conserva en memoria)
# -----------------------------------------------------------------------------

def summarize_result_for_group(res, index, path):
    """
    Extrae de un AnalysisResult solo lo necesario para el informe.

    Permite liberar el AnalysisResult y sus figuras inmediatamente despues de
    cada CSV, de forma que procesar 20 archivos no acumule memoria.
    """
    s = res.sampling
    m = res.motion
    fi = res.filter_info

    axis = {}
    for name in AXIS_NAMES:
        merged = dict(res.joint_metrics[name])
        merged.update(res.continuity_metrics[name])
        axis[name] = merged

    return {
        "index": index,
        "path": os.path.abspath(path),
        "name": os.path.basename(path),
        "ok": True,
        "error": None,
        "n_samples": s["n_samples"],
        "rows_original": res.clean_report["rows_original"],
        "duration_s": s["duration_s"],
        "freq_hz": s["freq_hz"],
        "dt_mean_ms": s["dt_mean_ms"],
        "dt_median_ms": s["dt_median_ms"],
        "dt_std_ms": s["dt_std_ms"],
        "dt_min_ms": s["dt_min_ms"],
        "dt_max_ms": s["dt_max_ms"],
        "motion_detected": bool(m["detected"]),
        "motion_reason": m["reason"],
        "t_start": m["t_start"] if m["detected"] else float("nan"),
        "t_end": m["t_end"] if m["detected"] else float("nan"),
        "motion_duration_s": m["duration_s"] if m["detected"] else 0.0,
        "velocity_threshold": m["effective_velocity_threshold"],
        "position_threshold": m["position_threshold"],
        "metric_window": res.metric_window_label(),
        "filter_applied": bool(fi.get("applied")),
        "filter_method": fi.get("method", "-"),
        "filter_params": fi.get("params", "-"),
        "filter_note": fi.get("note", ""),
        "processed_at": res.processed_at,
        "axis": axis,
        "figures": {},
        "skipped": [],
    }


def process_multiple_csvs(paths, tmp_dir, progress_cb=None):
    """
    Procesa la lista de CSV con EXACTAMENTE el mismo motor que el flujo
    individual.  Un archivo con error queda registrado y no detiene los demas.

    Devuelve la lista de resumenes (uno por archivo, en el orden seleccionado).
    """
    entries = []
    total = len(paths)
    for i, path in enumerate(paths, start=1):
        if progress_cb is not None:
            progress_cb(i, total, os.path.basename(path))

        figures = None
        try:
            res = analyze_file(path)                  # mismo motor de analisis
            figures, skipped = build_all_figures(res)  # mismas figuras
        except AnalysisError as exc:
            entries.append({"index": i, "path": os.path.abspath(path),
                            "name": os.path.basename(path), "ok": False,
                            "error": str(exc), "figures": {}, "skipped": []})
            continue
        except Exception as exc:  # noqa: BLE001 - un CSV roto no aborta el grupo
            entries.append({"index": i, "path": os.path.abspath(path),
                            "name": os.path.basename(path), "ok": False,
                            "error": "Error inesperado: %s" % exc,
                            "figures": {}, "skipped": []})
            continue

        try:
            entry = summarize_result_for_group(res, i, path)
            entry["skipped"] = skipped
            entry["figures"] = render_figures_to_tempdir(figures, tmp_dir, i)
            entries.append(entry)
        finally:
            # liberar figuras y resultado antes del siguiente CSV
            if figures:
                for fig in figures.values():
                    _release_figure(fig)
            figures = None
            res = None
    return entries


# -----------------------------------------------------------------------------
# Construccion del documento
# -----------------------------------------------------------------------------

def add_report_cover(document, paths, generated_at):
    """Portada limpia y tecnica + lista de archivos analizados."""
    _add_paragraph(document)
    _add_paragraph(document, "ANALISIS DE TELEMETRIA KUKA", bold=True,
                   size_pt=24, align=_DocxAlign.CENTER)
    _add_paragraph(document, "ESTUDIO COMPARATIVO KRL / MOVEIT2", bold=True,
                   size_pt=16, align=_DocxAlign.CENTER)
    _add_paragraph(document)
    _add_paragraph(document, "Informe de procesamiento multiple",
                   size_pt=13, align=_DocxAlign.CENTER)
    _add_paragraph(document)
    _add_paragraph(document, "Fecha de generacion: %s"
                   % generated_at.strftime("%d/%m/%Y %H:%M"),
                   size_pt=11, align=_DocxAlign.CENTER)
    _add_paragraph(document, "Cantidad de archivos procesados: %d" % len(paths),
                   size_pt=11, align=_DocxAlign.CENTER)
    _add_paragraph(document, "Herramienta: analisis_comparativo_gui.py v%s" % APP_VERSION,
                   size_pt=9, align=_DocxAlign.CENTER)
    _add_paragraph(document)

    _add_paragraph(document, "ARCHIVOS ANALIZADOS", bold=True, size_pt=12)
    for i, path in enumerate(paths, start=1):
        par = document.add_paragraph("%d. %s" % (i, os.path.basename(path)))
        par.paragraph_format.space_after = _DocxPt(2)
    _add_page_break(document)


def add_group_summary(document, entries):
    """1. RESUMEN GENERAL: tabla con una fila por CSV."""
    document.add_heading("1. RESUMEN GENERAL", level=1)
    document.add_heading("1.1 Tabla general de ensayos", level=2)

    headers = ["Ensayo", "Archivo", "Muestras", "Frecuencia [Hz]",
               "Periodo medio [ms]", "Periodo std [ms]", "Inicio mov. [s]",
               "Fin mov. [s]", "Duracion mov. [s]", "Movimiento"]
    rows = []
    for e in entries:
        if not e["ok"]:
            rows.append([e["index"], e["name"], "-", "-", "-", "-", "-", "-",
                         "-", "NO PROCESADO"])
            continue
        rows.append([
            e["index"], e["name"], e["n_samples"],
            _fmt(e["freq_hz"], 3), _fmt(e["dt_mean_ms"], 2),
            _fmt(e["dt_std_ms"], 2), _fmt(e["t_start"], 3),
            _fmt(e["t_end"], 3), _fmt(e["motion_duration_s"], 3),
            "Si" if e["motion_detected"] else "No",
        ])
    widths = [0.6, 3.0, 0.8, 0.85, 0.9, 0.85, 0.8, 0.75, 0.9, 1.05]
    _add_table(document, headers, rows, widths_in=widths, font_pt=8)
    _add_paragraph(document)


def add_group_statistics(document, entries):
    """1.2 estadisticos generales y 1.3 estadisticos articulares entre ensayos."""
    ok_entries = [e for e in entries if e["ok"]]

    document.add_heading("1.2 Estadisticos generales del grupo", level=2)
    if not ok_entries:
        _add_paragraph(document, "No hay ensayos procesados correctamente: no "
                                 "es posible calcular estadisticos del grupo.")
        _add_paragraph(document)
        return

    freqs = [e["freq_hz"] for e in ok_entries]
    durations_all = [e["duration_s"] for e in ok_entries]
    moving = [e for e in ok_entries if e["motion_detected"]]
    motion_durations = [e["motion_duration_s"] for e in moving]

    f = _mean_std_min_max(freqs)
    d = _mean_std_min_max(durations_all)
    md = _mean_std_min_max(motion_durations)

    rows = [
        ["Cantidad de ensayos procesados", len(ok_entries), "-"],
        ["Ensayos con movimiento detectado", len(moving), "-"],
        ["Ensayos sin movimiento detectado", len(ok_entries) - len(moving), "-"],
        ["Muestras totales analizadas", sum(e["n_samples"] for e in ok_entries), "-"],
        ["Frecuencia promedio", _fmt(f["mean"], 3), "Hz"],
        ["Frecuencia minima", _fmt(f["min"], 3), "Hz"],
        ["Frecuencia maxima", _fmt(f["max"], 3), "Hz"],
        ["Desviacion estandar de la frecuencia", _fmt(f["std"], 3), "Hz"],
        ["Duracion total registrada promedio", _fmt(d["mean"], 3), "s"],
        ["Duracion total registrada minima", _fmt(d["min"], 3), "s"],
        ["Duracion total registrada maxima", _fmt(d["max"], 3), "s"],
        ["Duracion de movimiento promedio *", _fmt(md["mean"], 3), "s"],
        ["Duracion de movimiento minima *", _fmt(md["min"], 3), "s"],
        ["Duracion de movimiento maxima *", _fmt(md["max"], 3), "s"],
        ["Desviacion estandar de la duracion de movimiento *", _fmt(md["std"], 3), "s"],
    ]
    _add_table(document, ["Parametro", "Valor", "Unidad"], rows,
               widths_in=[4.2, 1.8, 1.2])
    _add_paragraph(document,
                   "* Calculado unicamente sobre los %d ensayo(s) en los que se "
                   "detecto movimiento, para no sesgar la media con los ensayos "
                   "estaticos (duracion 0.000 s)." % len(moving), size_pt=8)
    _add_paragraph(document)

    document.add_heading("1.3 Estadisticos articulares entre ensayos", level=2)
    _add_paragraph(document,
                   "Cada valor resume, entre los %d ensayos procesados, la metrica "
                   "obtenida individualmente en cada CSV (el valor concreto de cada "
                   "ensayo aparece en la tabla de su propia seccion)."
                   % len(ok_entries), size_pt=9)

    for key, label, unit, decimals in GROUP_AXIS_METRICS:
        _add_paragraph(document, "%s [%s]" % (label, unit), bold=True, size_pt=10)
        rows = []
        for name in AXIS_NAMES:
            values = [e["axis"][name].get(key, float("nan")) for e in ok_entries]
            st = _mean_std_min_max(values)
            rows.append([name, _fmt(st["mean"], decimals), _fmt(st["std"], decimals),
                         _fmt(st["min"], decimals), _fmt(st["max"], decimals),
                         st["n"]])
        _add_table(document, ["Eje", "Media", "Desv. est.", "Minimo", "Maximo",
                              "N ensayos"], rows,
                   widths_in=[0.7, 1.6, 1.6, 1.6, 1.6, 1.0])
        _add_paragraph(document, size_pt=4)


def add_info_table(document, entry, section_number):
    """X.1 adquisicion y X.2 deteccion de movimiento."""
    document.add_heading("%d.1 Informacion de adquisicion" % section_number, level=2)
    rows = [
        ["Nombre del archivo", entry["name"], "-"],
        ["Numero de muestras validas", entry["n_samples"], "-"],
        ["Filas del CSV original", entry["rows_original"], "-"],
        ["Duracion total registrada", _fmt(entry["duration_s"], 3), "s"],
        ["Frecuencia efectiva", _fmt(entry["freq_hz"], 3), "Hz"],
        ["Periodo medio", _fmt(entry["dt_mean_ms"], 3), "ms"],
        ["Periodo mediano", _fmt(entry["dt_median_ms"], 3), "ms"],
        ["Desviacion estandar del periodo", _fmt(entry["dt_std_ms"], 3), "ms"],
        ["Periodo minimo", _fmt(entry["dt_min_ms"], 3), "ms"],
        ["Periodo maximo", _fmt(entry["dt_max_ms"], 3), "ms"],
    ]
    _add_table(document, ["Parametro", "Valor", "Unidad"], rows,
               widths_in=[3.6, 2.4, 1.0])
    _add_paragraph(document)

    document.add_heading("%d.2 Deteccion de movimiento y procesamiento"
                         % section_number, level=2)
    rows = [
        ["Movimiento detectado", "Si" if entry["motion_detected"] else "No", "-"],
        ["Inicio del movimiento", _fmt(entry["t_start"], 3), "s"],
        ["Fin del movimiento", _fmt(entry["t_end"], 3), "s"],
        ["Duracion del movimiento", _fmt(entry["motion_duration_s"], 3), "s"],
        ["Umbral de velocidad efectivo", _fmt(entry["velocity_threshold"], 6), "rad/s"],
        ["Umbral de recorrido", _fmt(entry["position_threshold"], 6), "rad"],
        ["Intervalo usado en las metricas", entry["metric_window"], "-"],
        ["Filtrado aplicado", "Si" if entry["filter_applied"] else "No", "-"],
        ["Metodo de filtrado", entry["filter_method"], "-"],
        ["Parametros del filtrado", entry["filter_params"], "-"],
        ["Metodo de derivacion",
         "numpy.gradient sobre tiempo real no uniforme", "-"],
    ]
    _add_table(document, ["Parametro", "Valor", "Unidad"], rows,
               widths_in=[2.6, 6.4, 0.9])
    if not entry["motion_detected"]:
        _add_paragraph(document,
                       "Sin movimiento detectado. Motivo: %s Las metricas "
                       "cinematicas de la tabla siguiente corresponden al ruido "
                       "residual del registro completo; no se atribuye a este "
                       "ensayo velocidad, aceleracion, jerk ni tiempo de "
                       "ejecucion significativos." % entry["motion_reason"],
                       size_pt=9)
    if entry["filter_note"]:
        _add_paragraph(document, "Observacion del filtrado: %s"
                       % entry["filter_note"], size_pt=9)
    _add_paragraph(document)


def add_metrics_table(document, entry, section_number):
    """X.3 tabla de metricas por articulacion (mismos valores que la GUI)."""
    document.add_heading("%d.3 Metricas articulares" % section_number, level=2)
    headers = ["Eje", "Vel. max. abs. [rad/s]", "Vel. RMS [rad/s]",
               "Acc. max. abs. [rad/s2]", "Acc. RMS [rad/s2]",
               "Jerk max. abs. [rad/s3]", "Jerk RMS [rad/s3]", "dq max. [rad]"]
    rows = []
    for name in AXIS_NAMES:
        a = entry["axis"][name]
        rows.append([
            name,
            _fmt(a["vel_max_abs_rad_s"], 5), _fmt(a["vel_rms_rad_s"], 5),
            _fmt(a["acc_max_abs_rad_s2"], 5), _fmt(a["acc_rms_rad_s2"], 5),
            _fmt(a["jerk_max_abs_rad_s3"], 5), _fmt(a["jerk_rms_rad_s3"], 5),
            _fmt(a["delta_q_max_rad"], 6),
        ])
    _add_table(document, headers, rows,
               widths_in=[0.6, 1.4, 1.25, 1.45, 1.3, 1.45, 1.3, 1.15], font_pt=8)
    _add_paragraph(document)

    document.add_heading("%d.4 Posicion articular (resumen numerico)"
                         % section_number, level=2)
    headers = ["Eje", "Inicial [deg]", "Final [deg]", "Minima [deg]",
               "Maxima [deg]", "Recorrido absoluto [deg]"]
    rows = []
    for name in AXIS_NAMES:
        a = entry["axis"][name]
        rows.append([name, _fmt(a["pos_inicial_deg"], 4), _fmt(a["pos_final_deg"], 4),
                     _fmt(a["pos_min_deg"], 4), _fmt(a["pos_max_deg"], 4),
                     _fmt(a["recorrido_abs_deg"], 4)])
    _add_table(document, headers, rows,
               widths_in=[0.7, 1.6, 1.6, 1.6, 1.6, 2.0])
    _add_paragraph(document)


def add_single_test_section(document, entry, section_number):
    """Seccion completa de un ensayo procesado correctamente."""
    document.add_heading("%d. ENSAYO %d" % (section_number, entry["index"]), level=1)
    _add_paragraph(document, "Archivo: %s" % entry["name"], bold=True, size_pt=11)
    _add_paragraph(document, entry["path"], size_pt=8)

    add_info_table(document, entry, section_number)
    add_metrics_table(document, entry, section_number)

    fig_number = 0
    sub_number = 4  # X.1 adquisicion, X.2 movimiento, X.3 metricas, X.4 posicion num.
    for key, sub_title, caption in GROUP_FIGURE_ORDER:
        sub_number += 1
        document.add_heading("%d.%d %s" % (section_number, sub_number, sub_title),
                             level=2)
        image_path = entry["figures"].get(key)
        if image_path is None or not os.path.isfile(image_path):
            reason = "datos no disponibles en el CSV"
            for fname, why in entry.get("skipped", []):
                if FIG_FILENAMES.get(key) == fname:
                    reason = why
            _add_paragraph(document,
                           "Grafica no generada: %s." % reason, size_pt=9)
            continue
        fig_number += 1
        add_matplotlib_figure_to_docx(document, image_path)
        add_figure_caption(document, "Figura %d.%d. %s del ensayo %d (%s)."
                           % (section_number, fig_number, caption,
                              entry["index"], entry["name"]))


def add_error_section(document, entry, section_number):
    """Seccion de un CSV que no pudo procesarse. No detiene el resto."""
    document.add_heading("%d. ENSAYO %d - NO PROCESADO"
                         % (section_number, entry["index"]), level=1)
    _add_paragraph(document, "Archivo: %s" % entry["name"], bold=True, size_pt=11)
    _add_paragraph(document, entry["path"], size_pt=8)
    _add_paragraph(document, "Estado: No procesado", bold=True)
    _add_paragraph(document, "Motivo:")
    for line in str(entry["error"]).splitlines():
        if line.strip():
            _add_paragraph(document, line.strip(), size_pt=9)
    _add_paragraph(document,
                   "El resto de los archivos del grupo se proceso con normalidad.",
                   size_pt=9)


def add_processing_summary(document, entries, paths, generated_at, section_number):
    """Seccion final: recuento de archivos y trazabilidad."""
    document.add_heading("%d. RESUMEN DE PROCESAMIENTO" % section_number, level=1)
    ok_entries = [e for e in entries if e["ok"]]
    err_entries = [e for e in entries if not e["ok"]]

    rows = [
        ["Archivos seleccionados", len(paths)],
        ["Procesados correctamente", len(ok_entries)],
        ["Con errores", len(err_entries)],
        ["Fecha de generacion", generated_at.strftime("%d/%m/%Y %H:%M:%S")],
        ["Herramienta", "analisis_comparativo_gui.py v%s" % APP_VERSION],
        ["Metodo de derivacion", "numpy.gradient(y, t), paso no uniforme"],
        ["SciPy disponible", "Si (v%s)" % SCIPY_VERSION if SCIPY_AVAILABLE else "No"],
        ["Resolucion de las figuras", "%d dpi" % REPORT_FIGURE_DPI],
    ]
    _add_table(document, ["Parametro", "Valor"], rows, widths_in=[3.6, 5.0])
    _add_paragraph(document)

    if err_entries:
        _add_paragraph(document, "Archivos con error", bold=True)
        _add_table(document, ["Ensayo", "Archivo", "Motivo"],
                   [[e["index"], e["name"],
                     " ".join(str(e["error"]).split())[:300]] for e in err_entries],
                   widths_in=[0.7, 2.8, 6.0], font_pt=8)
        _add_paragraph(document)

    _add_paragraph(document, "Limitaciones", bold=True)
    _add_paragraph(document,
                   "Indice completo de singularidad: no calculado. Requiere el "
                   "modelo cinematico / Jacobiano del KUKA KR6 R900 (parametros "
                   "DH o URDF); con las columnas del CSV no hay informacion "
                   "suficiente y no se implementa ninguna formula arbitraria.",
                   size_pt=9)
    _add_paragraph(document,
                   "Metricas que requieren medicion experimental externa y que "
                   "por tanto no se calculan ni se estiman a partir de la "
                   "telemetria: %s." % ", ".join(EXTERNAL_METRICS).lower(),
                   size_pt=9)


def save_group_report(document, path):
    """Escribe el .docx. Los errores de escritura se propagan como AnalysisError."""
    folder = os.path.dirname(path)
    try:
        os.makedirs(folder, exist_ok=True)
    except OSError as exc:
        raise AnalysisError("No se pudo crear la carpeta de informes:\n%s\n\n%s"
                            % (folder, exc))
    try:
        document.save(path)
    except PermissionError as exc:
        raise AnalysisError(
            "No se pudo escribir el documento (permiso denegado o archivo "
            "abierto en Word/LibreOffice):\n%s\n\n%s" % (path, exc))
    except OSError as exc:
        raise AnalysisError("No se pudo escribir el documento:\n%s\n\n%s"
                            % (path, exc))
    return path


def generate_group_report(paths, progress_cb=None):
    """
    Punto de entrada del informe multiple.

    Procesa cada CSV con el motor existente, construye UN unico .docx y borra
    los PNG temporales.  Devuelve (ruta_docx, entries, n_ok, n_error).
    """
    if not DOCX_AVAILABLE:
        raise AnalysisError(docx_requirement_message())
    if not paths:
        raise AnalysisError("No se selecciono ningun archivo CSV.")

    generated_at = datetime.datetime.now()
    tmp_dir = tempfile.mkdtemp(prefix="kuka_informe_grupal_")
    try:
        entries = process_multiple_csvs(paths, tmp_dir, progress_cb=progress_cb)

        if progress_cb is not None:
            progress_cb(len(paths), len(paths), "Construyendo el documento Word...")

        document = _new_report_document()
        add_report_cover(document, paths, generated_at)
        add_group_summary(document, entries)
        add_group_statistics(document, entries)

        for offset, entry in enumerate(entries):
            section_number = offset + 2          # la seccion 1 es el resumen general
            _add_page_break(document)
            if entry["ok"]:
                add_single_test_section(document, entry, section_number)
            else:
                add_error_section(document, entry, section_number)

        _add_page_break(document)
        add_processing_summary(document, entries, paths, generated_at,
                               len(entries) + 2)

        out_path = save_group_report(document, group_report_path(generated_at))
    finally:
        # los PNG temporales se eliminan siempre; los CSV originales y los
        # resultados individuales ya existentes no se tocan nunca
        shutil.rmtree(tmp_dir, ignore_errors=True)

    n_ok = sum(1 for e in entries if e["ok"])
    return out_path, entries, n_ok, len(entries) - n_ok


def open_document(path):
    """Abre el .docx con la aplicacion predeterminada del sistema, sin shell."""
    if not path or not os.path.isfile(path):
        raise AnalysisError("El documento no existe:\n%s" % path)
    try:
        if sys.platform.startswith("linux"):
            subprocess.Popen(["xdg-open", path],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        elif os.name == "nt":
            os.startfile(path)  # noqa: F821  (solo existe en Windows)
        else:
            raise AnalysisError("Plataforma no soportada para abrir documentos: %s"
                                % sys.platform)
    except FileNotFoundError:
        raise AnalysisError("No se encontro una aplicacion para abrir el "
                            "documento.\nArchivo:\n%s" % path)
    except OSError as exc:
        raise AnalysisError("No se pudo abrir el documento.\nDetalle: %s" % exc)


# =============================================================================
# 12. INTERFAZ GRAFICA
# =============================================================================

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk


PLOT_TABS = [
    ("position", "Posicion"),
    ("velocity", "Velocidad"),
    ("acceleration", "Aceleracion"),
    ("jerk", "Jerk"),
    ("continuity", "Continuidad"),
]

MAIN_TABLE_COLUMNS = [
    ("eje", "Eje", 55),
    ("vel_max", "Vel Max [rad/s]", 130),
    ("vel_rms", "Vel RMS [rad/s]", 130),
    ("acc_max", "Acc Max [rad/s2]", 135),
    ("acc_rms", "Acc RMS [rad/s2]", 135),
    ("jerk_max", "Jerk Max [rad/s3]", 140),
    ("jerk_rms", "Jerk RMS [rad/s3]", 140),
    ("dq_max", "dq Max [rad]", 125),
]

POS_TABLE_COLUMNS = [
    ("eje", "Eje", 55),
    ("ini", "Inicial [deg]", 120),
    ("fin", "Final [deg]", 120),
    ("min", "Minima [deg]", 120),
    ("max", "Maxima [deg]", 120),
    ("rec", "Recorrido abs. [deg]", 150),
]


class ComparativeAnalyzerApp(tk.Tk):
    """GUI del analizador comparativo."""

    def __init__(self):
        tk.Tk.__init__(self)
        self.title("%s  v%s" % (APP_TITLE, APP_VERSION))
        self.geometry("1280x820")
        self.minsize(1024, 680)

        self.csv_path = None
        self.result = None
        self.figures = {}
        self.skipped_figures = []
        self.summary_text = ""
        self.last_saved_folder = None
        self._canvases = []

        # --- estado de la reproduccion en RViz (funcion opcional) ------------
        self._rviz_playback_active = False
        self._rviz_playback_index = 0
        self._rviz_playback_start_monotonic = None
        self._rviz_after_id = None
        self._rviz_discovery_after_id = None
        self._rviz_discovery_deadline = None
        self._rviz_rclpy = None                  # modulo rclpy (lazy import)
        self._rviz_jointstate_cls = None         # clase JointState (lazy import)
        self._rviz_ros_node = None
        self._rviz_publisher = None
        self._rviz_rclpy_initialized_here = False
        self._rviz_messages_published = 0
        self._rviz_invalid_samples = 0

        # --- estado del informe grupal en Word (funcion aditiva) -------------
        self._group_running = False
        self._group_queue = None
        self._group_after_id = None
        self._group_saved_states = []
        self._group_report_path = None

        self._build_widgets()
        self._reset_state(clear_file=True)

    # -- construccion de la interfaz -----------------------------------------
    def _build_widgets(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        style.configure("Treeview.Heading", font=("TkDefaultFont", 9, "bold"))
        style.configure("Treeview", rowheight=24)

        # --- cabecera ---------------------------------------------------------
        header = ttk.Frame(self, padding=(10, 8, 10, 4))
        header.pack(fill="x")
        ttk.Label(header, text=APP_TITLE,
                  font=("TkDefaultFont", 14, "bold")).pack(anchor="w")
        ttk.Separator(header, orient="horizontal").pack(fill="x", pady=(6, 0))

        # --- seleccion de archivo --------------------------------------------
        top = ttk.Frame(self, padding=(10, 6))
        top.pack(fill="x")

        ttk.Label(top, text="Archivo seleccionado:").grid(row=0, column=0, sticky="w")
        self.file_var = tk.StringVar(value="(ninguno)")
        ttk.Label(top, textvariable=self.file_var, font=("TkDefaultFont", 10, "bold"),
                  foreground="#0d47a1").grid(row=0, column=1, sticky="w", padx=(8, 0))

        self.path_var = tk.StringVar(value="")
        ttk.Label(top, textvariable=self.path_var, foreground="#555555"
                  ).grid(row=1, column=0, columnspan=3, sticky="w", pady=(2, 6))

        buttons = ttk.Frame(top)
        buttons.grid(row=2, column=0, columnspan=3, sticky="w")
        self.btn_select = ttk.Button(buttons, text="Seleccionar CSV",
                                     command=self.on_select_csv)
        self.btn_select.pack(side="left")
        self.btn_process = ttk.Button(buttons, text="Procesar",
                                      command=self.on_process, state="disabled")
        self.btn_process.pack(side="left", padx=(8, 0))
        self.btn_clear = ttk.Button(buttons, text="Limpiar", command=self.on_clear)
        self.btn_clear.pack(side="left", padx=(8, 0))

        # Flujo independiente: varios CSV -> un unico documento Word.
        ttk.Separator(buttons, orient="vertical").pack(side="left", fill="y",
                                                       padx=(14, 14))
        self.btn_group_report = ttk.Button(
            buttons, text="Generar informe de varios CSV",
            command=self.on_generate_group_report)
        self.btn_group_report.pack(side="left")

        top.columnconfigure(2, weight=1)

        # la barra inferior se empaqueta antes que el cuaderno (side="bottom")
        self._build_bottom_bar()

        # --- cuaderno de pestanas ---------------------------------------------
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=(4, 4))

        self.tab_summary = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_summary, text="Resumen")
        self.summary_widget = self._make_text_view(self.tab_summary)
        self.summary_widget.configure(state="disabled")

        self.tab_metrics = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_metrics, text="Metricas")
        self._build_metrics_tab(self.tab_metrics)

        self.plot_frames = {}
        for key, label in PLOT_TABS:
            frame = ttk.Frame(self.notebook)
            self.notebook.add(frame, text=label)
            self.plot_frames[key] = frame

        self.tab_cartesian = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_cartesian, text="Cartesianas")
        self.cartesian_notebook = ttk.Notebook(self.tab_cartesian)
        self.cartesian_notebook.pack(fill="both", expand=True, padx=4, pady=4)
        frame = ttk.Frame(self.cartesian_notebook)
        self.cartesian_notebook.add(frame, text="Posicion XYZ")
        self.plot_frames["cart_pos"] = frame
        frame = ttk.Frame(self.cartesian_notebook)
        self.cartesian_notebook.add(frame, text="Orientacion ABC")
        self.plot_frames["cart_ori"] = frame

        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Muestreo")
        self.plot_frames["sampling"] = frame

        self.tab_limits = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_limits, text="Limitaciones")
        self._build_limits_tab(self.tab_limits)

    def _build_bottom_bar(self):
        """Barra inferior: estado y acciones. Se empaqueta ANTES del cuaderno
        para que nunca quede oculta por las graficas incrustadas."""
        bottom = ttk.Frame(self, padding=(10, 4, 10, 10))
        bottom.pack(fill="x", side="bottom")

        ttk.Label(bottom, text="Estado:").grid(row=0, column=0, sticky="w")
        self.status_var = tk.StringVar(value="Listo. Seleccione un CSV de telemetria.")
        self.status_label = ttk.Label(bottom, textvariable=self.status_var,
                                      font=("TkDefaultFont", 10, "bold"))
        self.status_label.grid(row=0, column=1, sticky="w", padx=(8, 0))

        self.saved_var = tk.StringVar(value="Ultima carpeta guardada: (ninguna)")
        ttk.Label(bottom, textvariable=self.saved_var, foreground="#555555"
                  ).grid(row=1, column=0, columnspan=4, sticky="w", pady=(2, 6))

        action = ttk.Frame(bottom)
        action.grid(row=2, column=0, columnspan=4, sticky="w")
        self.btn_plots = ttk.Button(action, text="Ver graficas",
                                    command=self.on_show_plots, state="disabled")
        self.btn_plots.pack(side="left")
        self.btn_rviz = ttk.Button(action, text="Probar con RViz",
                                   command=self.on_test_rviz, state="disabled")
        self.btn_rviz.pack(side="left", padx=(8, 0))
        self.btn_save = ttk.Button(action, text="Guardar resultados",
                                   command=self.on_save, state="disabled")
        self.btn_save.pack(side="left", padx=(8, 0))
        self.btn_open = ttk.Button(action, text="Abrir carpeta de resultados",
                                   command=self.on_open_folder, state="disabled")
        self.btn_open.pack(side="left", padx=(8, 0))
        self.btn_open_report = ttk.Button(action, text="Abrir informe",
                                          command=self.on_open_report,
                                          state="disabled")
        self.btn_open_report.pack(side="left", padx=(8, 0))

        # --- progreso del informe grupal (oculto mientras no se usa) ---------
        self.group_progress_frame = ttk.Frame(bottom)
        self.group_progress_frame.grid(row=3, column=0, columnspan=4,
                                       sticky="we", pady=(6, 0))
        self.group_progress_var = tk.StringVar(value="")
        ttk.Label(self.group_progress_frame,
                  textvariable=self.group_progress_var).pack(side="left")
        self.group_progress = ttk.Progressbar(self.group_progress_frame,
                                              orient="horizontal", length=300,
                                              mode="determinate")
        self.group_progress.pack(side="left", padx=(12, 0))
        self.group_progress_frame.grid_remove()

        self.report_var = tk.StringVar(value="Ultimo informe generado: (ninguno)")
        ttk.Label(bottom, textvariable=self.report_var, foreground="#555555"
                  ).grid(row=4, column=0, columnspan=4, sticky="w", pady=(6, 0))

        bottom.columnconfigure(3, weight=1)

    def _make_text_view(self, parent):
        """Text monoespaciado con barras vertical y horizontal."""
        frame = ttk.Frame(parent)
        frame.pack(fill="both", expand=True, padx=4, pady=4)
        widget = tk.Text(frame, wrap="none", font=("TkFixedFont", 10),
                         borderwidth=1, relief="sunken")
        vbar = ttk.Scrollbar(frame, orient="vertical", command=widget.yview)
        hbar = ttk.Scrollbar(frame, orient="horizontal", command=widget.xview)
        widget.configure(yscrollcommand=vbar.set, xscrollcommand=hbar.set)
        widget.grid(row=0, column=0, sticky="nsew")
        vbar.grid(row=0, column=1, sticky="ns")
        hbar.grid(row=1, column=0, sticky="ew")
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        return widget

    def _make_scrollable(self, parent):
        """Marco desplazable verticalmente (las tablas no caben siempre)."""
        canvas = tk.Canvas(parent, highlightthickness=0, borderwidth=0)
        vbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        inner = ttk.Frame(canvas)
        window = canvas.create_window((0, 0), window=inner, anchor="nw")
        inner.bind("<Configure>",
                   lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>",
                    lambda e: canvas.itemconfigure(window, width=e.width))

        def _wheel(event):
            if event.num == 4:
                canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                canvas.yview_scroll(1, "units")
            else:
                canvas.yview_scroll(-1 * int(event.delta / 120), "units")

        def _bind(_event):
            canvas.bind_all("<MouseWheel>", _wheel)
            canvas.bind_all("<Button-4>", _wheel)
            canvas.bind_all("<Button-5>", _wheel)

        def _unbind(_event):
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")

        canvas.bind("<Enter>", _bind)
        canvas.bind("<Leave>", _unbind)
        return inner

    def _build_metrics_tab(self, parent):
        outer = ttk.Frame(parent, padding=6)
        outer.pack(fill="both", expand=True)
        container = self._make_scrollable(outer)

        self.window_var = tk.StringVar(value="Intervalo de metricas: -")
        ttk.Label(container, textvariable=self.window_var,
                  font=("TkDefaultFont", 10, "bold")).pack(anchor="w", pady=(0, 4))

        ttk.Label(container, text="Cinematica articular (intervalo de movimiento)",
                  font=("TkDefaultFont", 10, "bold")).pack(anchor="w")
        self.tree_main = self._make_tree(container, MAIN_TABLE_COLUMNS, height=6)

        ttk.Label(container, text="Posicion articular",
                  font=("TkDefaultFont", 10, "bold")).pack(anchor="w", pady=(10, 0))
        self.tree_pos = self._make_tree(container, POS_TABLE_COLUMNS, height=6)

        ttk.Label(container, text="Adquisicion / muestreo",
                  font=("TkDefaultFont", 10, "bold")).pack(anchor="w", pady=(10, 0))
        cols = [("param", "Parametro", 320), ("valor", "Valor", 160),
                ("unidad", "Unidad", 100)]
        self.tree_sampling = self._make_tree(container, cols, height=13)

    def _make_tree(self, parent, columns, height):
        keys = [c[0] for c in columns]
        tree = ttk.Treeview(parent, columns=keys, show="headings", height=height)
        for key, text, width in columns:
            tree.heading(key, text=text)
            tree.column(key, width=width, anchor="center", stretch=True)
        tree.pack(fill="x", pady=(2, 0))
        return tree

    def _build_limits_tab(self, parent):
        widget = self._make_text_view(parent)
        lines = [
            "LIMITACIONES DEL ANALISIS BASADO UNICAMENTE EN LA TELEMETRIA CSV",
            "=" * 72,
            "",
            "Indice completo de singularidad:",
            "    No calculado.",
            "",
            "    Se requiere el modelo cinematico / Jacobiano del KUKA KR6 R900",
            "    (parametros DH, URDF, dimensiones geometricas).  Las columnas",
            "    disponibles en el CSV (posiciones articulares y pose registrada)",
            "    no aportan informacion suficiente, y esta herramienta no",
            "    implementa ninguna formula arbitraria de singularidad.",
            "",
            "Metricas que requieren medicion experimental externa:",
        ]
        for item in EXTERNAL_METRICS:
            lines.append("    - %s : requiere medicion experimental externa." % item)
        lines += [
            "",
            "    Estas magnitudes no se deducen de la telemetria y no se generan",
            "    valores aproximados ni estimados.",
            "",
            "Orientacion cartesiana (A, B, C):",
            "    Se representa unicamente su evolucion temporal registrada.",
            "    No se calcula una velocidad angular del efector derivando A, B y C",
            "    por separado, porque esa operacion no es valida sin la",
            "    transformacion cinematica correcta de la representacion.",
            "",
            "Estas mismas limitaciones quedan registradas en resumen.txt.",
        ]
        widget.insert("1.0", "\n".join(lines))
        widget.configure(state="disabled")

    # -- gestion de estado ----------------------------------------------------
    def _clear_canvases(self):
        for canvas in self._canvases:
            try:
                canvas.get_tk_widget().destroy()
            except tk.TclError:
                pass
        self._canvases = []
        for frame in self.plot_frames.values():
            for child in frame.winfo_children():
                child.destroy()
        self.figures = {}

    def _reset_state(self, clear_file=False):
        # Cambiar de CSV o limpiar detiene cualquier reproduccion en curso: no
        # deben seguir saliendo muestras del ensayo anterior.
        self._cancel_rviz_playback()
        self.result = None
        self.skipped_figures = []
        self.summary_text = ""
        self._clear_canvases()

        self.summary_widget.configure(state="normal")
        self.summary_widget.delete("1.0", "end")
        self.summary_widget.configure(state="disabled")

        for tree in (self.tree_main, self.tree_pos, self.tree_sampling):
            for item in tree.get_children():
                tree.delete(item)
        self.window_var.set("Intervalo de metricas: -")

        self.btn_save.configure(state="disabled")
        self.btn_plots.configure(state="disabled")
        self.btn_rviz.configure(state="disabled")
        if clear_file:
            self.csv_path = None
            self.file_var.set("(ninguno)")
            self.path_var.set("")
            self.btn_process.configure(state="disabled")
            self.btn_open.configure(state="disabled")
            self.last_saved_folder = None
            self.saved_var.set("Ultima carpeta guardada: (ninguna)")

    def _set_status(self, text, color="#1b5e20"):
        self.status_var.set(text)
        try:
            self.status_label.configure(foreground=color)
        except tk.TclError:
            pass
        self.update_idletasks()

    # -- acciones -------------------------------------------------------------
    def on_select_csv(self):
        initial = os.path.join(os.path.dirname(SCRIPT_DIR), "logs")
        if not os.path.isdir(initial):
            initial = os.path.dirname(SCRIPT_DIR)
        path = filedialog.askopenfilename(
            title="Seleccionar CSV de telemetria KUKA",
            initialdir=initial,
            filetypes=[("CSV de telemetria", "kuka_telemetry_*.csv"),
                       ("Archivos CSV", "*.csv"),
                       ("Todos los archivos", "*.*")])
        if not path:
            return
        # cambiar de CSV limpia por completo el estado anterior
        self._reset_state(clear_file=True)
        self.csv_path = path
        self.file_var.set(os.path.basename(path))
        self.path_var.set(path)
        self.btn_process.configure(state="normal")
        self._set_status("Archivo seleccionado. Pulse 'Procesar'.", "#0d47a1")

    def on_process(self):
        if not self.csv_path:
            messagebox.showwarning("Sin archivo",
                                   "Primero seleccione un archivo CSV.")
            return
        self._reset_state(clear_file=False)
        self._set_status("Procesando...", "#ef6c00")
        self.configure(cursor="watch")
        self.update_idletasks()
        try:
            result = analyze_file(self.csv_path)
            figures, skipped = build_all_figures(result)
        except AnalysisError as exc:
            self.configure(cursor="")
            self._set_status("Error: no se pudo procesar el archivo.", "#b71c1c")
            messagebox.showerror("No se pudo procesar el CSV", str(exc))
            return
        except Exception as exc:  # noqa: BLE001 - la GUI nunca debe cerrarse
            self.configure(cursor="")
            self._set_status("Error inesperado durante el procesamiento.", "#b71c1c")
            messagebox.showerror(
                "Error inesperado",
                "Ocurrio un error inesperado al procesar el archivo:\n\n%s\n\n%s"
                % (exc, traceback.format_exc(limit=3)))
            return

        self.result = result
        self.figures = figures
        self.skipped_figures = skipped
        self.summary_text = create_summary_txt(result, skipped)

        try:
            self._fill_summary()
            self._fill_tables()
            self._embed_figures()
        except Exception as exc:  # noqa: BLE001
            self.configure(cursor="")
            self._set_status("Error al mostrar los resultados.", "#b71c1c")
            messagebox.showerror("Error al mostrar resultados",
                                 "%s\n\n%s" % (exc, traceback.format_exc(limit=3)))
            return

        self.configure(cursor="")
        self.btn_save.configure(state="normal")
        self.btn_plots.configure(state="normal")
        self.btn_rviz.configure(state="normal")

        if result.motion["detected"]:
            self._set_status("Procesamiento completado correctamente.", "#1b5e20")
        else:
            self._set_status("Procesado. SIN MOVIMIENTO DETECTADO.", "#ef6c00")

        if result.warnings:
            messagebox.showwarning(
                "Avisos del procesamiento",
                "El archivo se proceso correctamente, con estos avisos:\n\n- %s"
                % "\n- ".join(result.warnings[:12]))

    def on_clear(self):
        self._reset_state(clear_file=True)
        self._set_status("Listo. Seleccione un CSV de telemetria.", "#1b5e20")

    def on_show_plots(self):
        if self.result is None:
            return
        self.notebook.select(self.plot_frames["position"])

    # -- reproduccion en RViz (funcion OPCIONAL, solo visualizacion) ----------
    #
    #   CSV ya procesado -> self.result.q_rad [rad] -> sensor_msgs/JointState
    #   -> /fake_joint_states -> joint_state_publisher -> /joint_states -> RViz
    #
    # Esta funcion NO controla el KUKA: no abre sockets TCP, no habla
    # EthernetKRL / EKI, no publica target_json ni ningun comando, no usa
    # MoveIt planning, ActionClients ni FollowJointTrajectory.  Es OFFLINE: no
    # necesita el robot conectado, solo el CSV grabado y el visualizador.

    def _ensure_rviz_publisher(self):
        """
        Prepara ROS2 la primera vez que se pulsa el boton (lazy import).

        El import de rclpy / sensor_msgs se hace AQUI y no al principio del
        archivo, para que la GUI arranque y el analisis funcione en entornos
        sin ROS2.  Devuelve True si el publicador esta listo.
        """
        if self._rviz_publisher is not None:
            return True

        try:
            import rclpy
            from rclpy.qos import (
                DurabilityPolicy,
                HistoryPolicy,
                QoSProfile,
                ReliabilityPolicy,
            )
            from sensor_msgs.msg import JointState
        except Exception as exc:  # noqa: BLE001 - la GUI nunca debe cerrarse
            self._set_status("ROS2 no disponible: el analisis sigue funcionando.",
                             "#ef6c00")
            messagebox.showerror(
                "ROS2 no disponible",
                "%s\n\nDetalle: %s" % (ROS2_UNAVAILABLE_MESSAGE, exc))
            return False

        try:
            if not rclpy.ok():
                rclpy.init(args=None)
                # Solo se apagara rclpy al cerrar si lo inicializo este script.
                self._rviz_rclpy_initialized_here = True
            node = rclpy.create_node(RVIZ_REPLAY_NODE_NAME)
            # QoS identico al del rviz_mirror_node ya validado.
            qos = QoSProfile(
                reliability=ReliabilityPolicy.RELIABLE,
                durability=DurabilityPolicy.VOLATILE,
                history=HistoryPolicy.KEEP_LAST,
                depth=RVIZ_REPLAY_QOS_DEPTH,
            )
            # UNICO publicador de esta herramienta.
            publisher = node.create_publisher(JointState, RVIZ_REPLAY_TOPIC, qos)
        except Exception as exc:  # noqa: BLE001
            self._set_status("No se pudo inicializar ROS2.", "#b71c1c")
            messagebox.showerror(
                "Error al inicializar ROS2",
                "No se pudo crear el nodo o el publicador ROS2.\n\n%s\n\n%s"
                % (exc, traceback.format_exc(limit=3)))
            return False

        self._rviz_rclpy = rclpy
        self._rviz_jointstate_cls = JointState
        self._rviz_ros_node = node
        self._rviz_publisher = publisher
        return True

    def _validate_playback_data(self):
        """Comprueba que el resultado procesado sirve para reproducir."""
        res = self.result
        if res is None:
            return "No hay ningun CSV procesado."
        if res.t is None or res.q_rad is None:
            return "El resultado procesado no contiene tiempo o posiciones."
        t = np.asarray(res.t, dtype=float)
        q = np.asarray(res.q_rad, dtype=float)
        if t.ndim != 1 or q.ndim != 2:
            return "Las series de tiempo y posicion no tienen la forma esperada."
        if q.shape[1] != 6:
            return ("Cada muestra debe contener exactamente 6 articulaciones "
                    "(se encontraron %d)." % q.shape[1])
        if t.shape[0] != q.shape[0]:
            return ("El numero de tiempos (%d) y de muestras articulares (%d) "
                    "no coincide." % (t.shape[0], q.shape[0]))
        if t.shape[0] < 1:
            return "No hay ninguna muestra que reproducir."
        return None

    def _publish_rviz_sample(self, index):
        """
        Publica UNA muestra en /fake_joint_states.

        Usa res.q_rad, que es la posicion articular CRUDA registrada convertida
        de grados a radianes.  NO se usa res.q_used_rad (puede estar filtrada
        para el analisis): se reproduce exactamente lo registrado, sin filtrar,
        sin interpolar y sin suavizar.  Devuelve True si se publico.
        """
        q = np.asarray(self.result.q_rad[index], dtype=float)
        if q.shape[0] != 6 or not np.all(np.isfinite(q)):
            # Muestra invalida (NaN / Inf): se descarta, pero la reproduccion
            # del resto del ensayo continua.
            self._rviz_invalid_samples += 1
            return False

        msg = self._rviz_jointstate_cls()
        # Tiempo de REPRODUCCION ROS2, no el timestamp original del KUKA. El
        # timestamp registrado solo decide CUANDO se publica cada muestra.
        msg.header.stamp = self._rviz_ros_node.get_clock().now().to_msg()
        msg.name = list(RVIZ_REPLAY_JOINT_NAMES)
        msg.position = [float(v) for v in q]
        msg.velocity = []
        msg.effort = []
        self._rviz_publisher.publish(msg)
        self._rviz_messages_published += 1
        return True

    def _schedule_rviz_sample(self):
        """
        Programa la siguiente muestra SIN acumular drift.

        Cada muestra se compara contra el instante ABSOLUTO de reproduccion
        (inicio_monotonico + t[i]), no contra el dt de la muestra anterior: un
        retraso puntual de Tkinter se absorbe en la siguiente espera en vez de
        arrastrarse hasta el final.  Asi un ensayo de 59.14 s dura ~59.14 s.
        """
        index = self._rviz_playback_index
        target = self._rviz_playback_start_monotonic + float(self.result.t[index])
        remaining = target - time.monotonic()
        delay_ms = int(max(0.0, remaining) * 1000.0)
        self._rviz_after_id = self.after(delay_ms, self._rviz_playback_step)

    def _rviz_playback_step(self):
        """Publica la muestra actual y programa la siguiente. No bloquea."""
        self._rviz_after_id = None
        if not self._rviz_playback_active or self.result is None:
            return

        total = int(np.asarray(self.result.t).shape[0])
        index = self._rviz_playback_index
        if index >= total:
            self._finish_rviz_playback()
            return

        try:
            self._publish_rviz_sample(index)
        except Exception as exc:  # noqa: BLE001
            self._cancel_rviz_playback()
            self._set_status("Error durante la reproduccion en RViz.", "#b71c1c")
            if self.result is not None:
                self.btn_rviz.configure(state="normal")
            messagebox.showerror(
                "Error durante la reproduccion",
                "Fallo al publicar la muestra %d de %d:\n\n%s"
                % (index + 1, total, exc))
            return

        self._rviz_playback_index = index + 1

        if (self._rviz_playback_index % RVIZ_REPLAY_PROGRESS_EVERY == 0
                and self._rviz_playback_index < total):
            self._set_status("Reproduciendo en RViz: %d / %d"
                             % (self._rviz_playback_index, total), "#ef6c00")

        if self._rviz_playback_index >= total:
            self._finish_rviz_playback()
        else:
            self._schedule_rviz_sample()

    def _cancel_rviz_playback(self):
        """
        Cancela la reproduccion Y la espera de descubrimiento pendientes.

        Seguro de llamar siempre.  No deja callbacks huerfanos de Tk.after ni
        de la fase de discovery ni de la linea temporal.
        """
        was_pending = False
        for attr in ("_rviz_discovery_after_id", "_rviz_after_id"):
            after_id = getattr(self, attr, None)
            if after_id is not None:
                was_pending = True
                try:
                    self.after_cancel(after_id)
                except (tk.TclError, ValueError):
                    pass
            setattr(self, attr, None)
        self._rviz_discovery_deadline = None
        was_active = bool(getattr(self, "_rviz_playback_active", False))
        self._rviz_playback_active = False
        self._rviz_playback_index = 0
        return was_active or was_pending

    def _finish_rviz_playback(self):
        """Cierra la reproduccion e informa del resultado."""
        self._rviz_playback_active = False
        self._rviz_after_id = None

        elapsed = 0.0
        if self._rviz_playback_start_monotonic is not None:
            elapsed = time.monotonic() - self._rviz_playback_start_monotonic
        recorded = 0.0
        if self.result is not None and self.result.t is not None:
            t = np.asarray(self.result.t, dtype=float)
            if t.shape[0]:
                recorded = float(t[-1] - t[0])

        if self.result is not None:
            self.btn_rviz.configure(state="normal")
        self._set_status("Reproduccion RViz completada.", "#1b5e20")

        detail = (
            "Reproduccion RViz completada.\n\n"
            "Muestras publicadas:      %d\n"
            "Muestras invalidas:       %d\n"
            "Duracion reproducida:     %.3f s\n"
            "Duracion registrada:      %.3f s\n\n"
            "Topico: %s\nTipo:   sensor_msgs/msg/JointState"
            % (self._rviz_messages_published, self._rviz_invalid_samples,
               elapsed, recorded, RVIZ_REPLAY_TOPIC))
        if self._rviz_invalid_samples:
            detail += ("\n\nLas muestras invalidas (NaN / Inf) no se publicaron; "
                       "el resto del ensayo si se reprodujo.")
        messagebox.showinfo("Reproduccion RViz completada", detail)

    def on_test_rviz(self):
        """Reproduce en RViz el registro COMPLETO del CSV ya procesado."""
        if self.result is None:
            messagebox.showwarning(
                "Sin datos",
                "Primero seleccione y procese un archivo CSV.")
            return
        if self._rviz_playback_active or self._rviz_discovery_after_id is not None:
            return

        problem = self._validate_playback_data()
        if problem is not None:
            messagebox.showerror("No se puede reproducir", problem)
            return

        if not self._ensure_rviz_publisher():
            return

        # El boton queda deshabilitado durante el discovery Y durante toda la
        # reproduccion; se reactiva en _finish_rviz_playback.
        self.btn_rviz.configure(state="disabled")
        self._set_status("Preparando reproduccion RViz...", "#ef6c00")
        self._rviz_discovery_deadline = time.monotonic() + RVIZ_DISCOVERY_TIMEOUT_S
        self._wait_for_rviz_subscriber()

    def _wait_for_rviz_subscriber(self):
        """
        Espera NO bloqueante a que DDS empareje al subscriber, con timeout.

        Sin esto, la muestra 0 puede publicarse antes de que el subscriber de
        /fake_joint_states se haya descubierto y perderse.  Se sondea
        get_subscription_count() cada RVIZ_DISCOVERY_POLL_MS con Tk.after (no
        se usa time.sleep: no debe bloquearse Tkinter).

        Este tiempo de espera NO forma parte de la reproduccion: el reloj de la
        linea temporal se arranca despues, en _start_rviz_playback().
        """
        self._rviz_discovery_after_id = None

        # Si mientras esperabamos se limpio el estado o se cambio de CSV, no
        # hay nada que reproducir: _reset_state ya dejo el boton deshabilitado.
        if self.result is None or self._rviz_publisher is None:
            self._rviz_discovery_deadline = None
            return

        try:
            subscriber_count = self._rviz_publisher.get_subscription_count()
        except Exception:  # noqa: BLE001 - la GUI nunca debe cerrarse
            subscriber_count = 0

        if subscriber_count >= 1:
            self._start_rviz_playback()
            return

        if (self._rviz_discovery_deadline is None
                or time.monotonic() >= self._rviz_discovery_deadline):
            messagebox.showwarning(
                "Sin suscriptores",
                "No se detectó ningún suscriptor en /fake_joint_states.\n"
                "La reproducción comenzará igualmente.")
            self._start_rviz_playback()
            return

        self._rviz_discovery_after_id = self.after(
            RVIZ_DISCOVERY_POLL_MS, self._wait_for_rviz_subscriber)

    def _start_rviz_playback(self):
        """
        Arranca la linea temporal, una vez terminado el discovery.

        El reloj de reproduccion se fija AQUI y no al pulsar el boton, para que
        lo que haya tardado DDS (0 ms, 200 ms, 500 ms...) no desplace el
        ensayo: un CSV de 59.14 s dura ~59.14 s en cualquier caso.
        """
        self._rviz_discovery_deadline = None
        if self.result is None or self._rviz_publisher is None:
            return

        # Se reproduce el REGISTRO COMPLETO (indices 0..N-1), no analysis_slice
        # ni la ventana de movimiento: asi se ve tambien el tiempo quieto antes
        # y despues del movimiento.
        total = int(np.asarray(self.result.t).shape[0])
        self._rviz_playback_active = True
        self._rviz_playback_index = 0
        self._rviz_messages_published = 0
        self._rviz_invalid_samples = 0
        self.btn_rviz.configure(state="disabled")
        self._set_status("Reproduciendo CSV en RViz...", "#ef6c00")

        self._rviz_playback_start_monotonic = time.monotonic()
        # La primera muestra sale inmediatamente; las siguientes se programan
        # segun los tiempos reales registrados en self.result.t.
        self._rviz_playback_step()
        if total <= 0:
            self._finish_rviz_playback()

    def _shutdown_rviz_ros(self):
        """Destruye el nodo y apaga rclpy solo si lo inicializo este script."""
        node = self._rviz_ros_node
        self._rviz_publisher = None
        self._rviz_ros_node = None
        if node is not None:
            try:
                node.destroy_node()
            except Exception:  # noqa: BLE001
                pass
        if self._rviz_rclpy_initialized_here and self._rviz_rclpy is not None:
            try:
                self._rviz_rclpy.shutdown()
            except Exception:  # noqa: BLE001
                pass
        self._rviz_rclpy_initialized_here = False

    def on_close(self):
        """Cierre limpio: cancelar playback, cerrar ROS2 y destruir la ventana."""
        if self._group_after_id is not None:
            try:
                self.after_cancel(self._group_after_id)
            except tk.TclError:
                pass
            self._group_after_id = None
        self._cancel_rviz_playback()
        self._shutdown_rviz_ros()
        try:
            self.destroy()
        except tk.TclError:
            pass

    def on_save(self):
        if self.result is None:
            messagebox.showwarning("Sin resultados",
                                   "Procese un CSV antes de guardar.")
            return
        self._set_status("Guardando resultados...", "#ef6c00")
        self.configure(cursor="watch")
        self.update_idletasks()
        try:
            folder, written = save_results(self.result, self.figures,
                                           self.skipped_figures)
        except AnalysisError as exc:
            self.configure(cursor="")
            self._set_status("Error al guardar. Los resultados siguen en memoria.",
                             "#b71c1c")
            messagebox.showerror("Error al guardar", str(exc))
            return
        except Exception as exc:  # noqa: BLE001
            self.configure(cursor="")
            self._set_status("Error al guardar. Los resultados siguen en memoria.",
                             "#b71c1c")
            messagebox.showerror(
                "Error al guardar",
                "No se pudieron escribir los resultados:\n\n%s\n\n%s"
                % (exc, traceback.format_exc(limit=3)))
            return

        self.configure(cursor="")
        self.last_saved_folder = folder
        self.saved_var.set("Ultima carpeta guardada: %s" % folder)
        self.btn_open.configure(state="normal")
        self._set_status("Resultados guardados correctamente.", "#1b5e20")

        detail = "\n".join("  - %s" % w for w in written)
        extra = ""
        if self.skipped_figures:
            extra = "\n\nGraficas no generadas:\n" + "\n".join(
                "  - %s (%s)" % (f, r) for f, r in self.skipped_figures)
        messagebox.showinfo(
            "Guardado completado",
            "Resultados guardados correctamente.\n\nCarpeta:\n%s\n\nArchivos:\n%s%s"
            % (folder, detail, extra))

    # -- informe grupal en Word (funcion ADITIVA, flujo independiente) --------
    #
    #   Seleccionar varios CSV -> hilo de trabajo -> un unico .docx
    #
    #   El hilo NUNCA toca widgets: se comunica por cola y todas las
    #   actualizaciones de la GUI ocurren en _poll_group_queue(), que corre en
    #   el hilo principal mediante after().
    # -------------------------------------------------------------------------
    def _set_group_busy(self, busy, total=0):
        """
        Bloquea/restaura los botones durante la generacion del informe.

        Guarda el estado exacto previo de cada boton y lo restaura al terminar,
        de modo que el flujo individual queda igual que antes.
        """
        widgets = [self.btn_select, self.btn_process, self.btn_clear,
                   self.btn_group_report, self.btn_plots, self.btn_rviz,
                   self.btn_save, self.btn_open, self.btn_open_report]
        if busy:
            self._group_saved_states = []
            for widget in widgets:
                try:
                    self._group_saved_states.append((widget, str(widget["state"])))
                    widget.configure(state="disabled")
                except tk.TclError:
                    pass
            self.group_progress.configure(maximum=max(int(total), 1), value=0)
            self.group_progress_var.set("Preparando...")
            self.group_progress_frame.grid()
        else:
            for widget, state in self._group_saved_states:
                try:
                    widget.configure(state=state)
                except tk.TclError:
                    pass
            self._group_saved_states = []
            self.group_progress_var.set("")
            self.group_progress_frame.grid_remove()

    def on_generate_group_report(self):
        """Selecciona varios CSV y lanza la generacion del informe en Word."""
        if self._group_running:
            messagebox.showinfo("Informe en curso",
                                "Ya hay un informe generandose. Espere a que "
                                "termine.")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Falta la libreria python-docx",
                                 docx_requirement_message())
            return

        initial = os.path.join(os.path.dirname(SCRIPT_DIR), "logs")
        if not os.path.isdir(initial):
            initial = os.path.dirname(SCRIPT_DIR)
        selected = filedialog.askopenfilenames(
            title="Seleccionar los CSV del grupo (seleccion multiple)",
            initialdir=initial,
            filetypes=[("Archivos CSV", "*.csv"),
                       ("CSV de telemetria", "kuka_telemetry_*.csv"),
                       ("Todos los archivos", "*.*")])
        if not selected:
            return

        paths = list(selected)
        self._group_running = True
        self._group_report_path = None
        self._group_queue = _queue.Queue()
        self._set_group_busy(True, len(paths))
        self._set_status("Generando informe de %d archivos..." % len(paths),
                         "#ef6c00")

        worker = threading.Thread(target=self._group_worker,
                                  args=(paths, self._group_queue), daemon=True)
        worker.start()
        self._group_after_id = self.after(150, self._poll_group_queue)

    def _group_worker(self, paths, out_queue):
        """Hilo secundario: solo calcula y publica mensajes en la cola."""
        def progress_cb(index, total, name):
            out_queue.put(("progress", index, total, name))

        try:
            out_path, entries, n_ok, n_err = generate_group_report(
                paths, progress_cb=progress_cb)
            out_queue.put(("done", out_path, n_ok, n_err, len(paths)))
        except AnalysisError as exc:
            out_queue.put(("error", str(exc)))
        except Exception as exc:  # noqa: BLE001 - la GUI nunca debe cerrarse
            out_queue.put(("error",
                           "Error inesperado al generar el informe:\n\n%s\n\n%s"
                           % (exc, traceback.format_exc(limit=3))))

    def _poll_group_queue(self):
        """Hilo principal: consume la cola y actualiza los widgets."""
        self._group_after_id = None
        finished = False
        try:
            while True:
                message = self._group_queue.get_nowait()
                kind = message[0]

                if kind == "progress":
                    _, index, total, name = message
                    self.group_progress.configure(maximum=max(int(total), 1),
                                                  value=int(index))
                    if str(name).lower().endswith(".csv"):
                        text = "Procesando archivo %d de %d: %s" % (index, total, name)
                    else:
                        text = str(name)
                    self.group_progress_var.set(text)
                    self._set_status(text, "#ef6c00")

                elif kind == "done":
                    _, out_path, n_ok, n_err, n_selected = message
                    finished = True
                    self._group_running = False
                    self._set_group_busy(False)
                    self._group_report_path = out_path
                    self.btn_open_report.configure(state="normal")
                    self.report_var.set("Ultimo informe generado: %s" % out_path)
                    if n_err:
                        self._set_status(
                            "Informe generado con %d archivo(s) con error." % n_err,
                            "#ef6c00")
                        messagebox.showwarning(
                            "Informe generado",
                            "Archivos seleccionados: %d\nProcesados correctamente: "
                            "%d\nErrores: %d\n\nEl informe fue generado incluyendo "
                            "el detalle del archivo con error.\n\nDocumento:\n%s"
                            % (n_selected, n_ok, n_err, out_path))
                    else:
                        self._set_status("Informe generado correctamente.", "#1b5e20")
                        messagebox.showinfo(
                            "Informe generado",
                            "Informe generado correctamente.\n\nArchivos "
                            "seleccionados: %d\nProcesados correctamente: %d\n"
                            "Errores: 0\n\nDocumento:\n%s"
                            % (n_selected, n_ok, out_path))
                    break

                elif kind == "error":
                    _, detail = message
                    finished = True
                    self._group_running = False
                    self._set_group_busy(False)
                    self._set_status("Error al generar el informe.", "#b71c1c")
                    messagebox.showerror("No se pudo generar el informe", detail)
                    break
        except _queue.Empty:
            pass

        if self._group_running and not finished:
            self._group_after_id = self.after(150, self._poll_group_queue)

    def on_open_report(self):
        """Abre el ultimo informe generado con la aplicacion del sistema."""
        if not self._group_report_path:
            return
        try:
            open_document(self._group_report_path)
        except AnalysisError as exc:
            messagebox.showerror("No se pudo abrir el informe", str(exc))

    def on_open_folder(self):
        if not self.last_saved_folder:
            return
        try:
            open_results_folder(self.last_saved_folder)
        except AnalysisError as exc:
            messagebox.showerror("No se pudo abrir la carpeta", str(exc))

    # -- volcado de resultados en la GUI --------------------------------------
    def _fill_summary(self):
        res = self.result
        s = res.sampling
        m = res.motion
        head = [
            "=" * 72,
            "     %s" % APP_TITLE,
            "=" * 72,
            "",
            "Archivo seleccionado:",
            "  %s" % res.csv_name,
            "",
            "-" * 72,
            "RESUMEN DEL ENSAYO",
            "-" * 72,
            "Muestras:                   %d" % s["n_samples"],
            "Frecuencia efectiva:        %s Hz" % _fmt(s["freq_hz"], 3),
            "Periodo medio:              %s ms" % _fmt(s["dt_mean_ms"], 2),
            "Duracion registrada:        %s s" % _fmt(s["duration_s"], 2),
        ]
        if m["detected"]:
            head += [
                "Duracion del movimiento:    %s s" % _fmt(m["duration_s"], 2),
                "Inicio movimiento:          %s s" % _fmt(m["t_start"], 2),
                "Fin movimiento:             %s s" % _fmt(m["t_end"], 2),
            ]
        else:
            head += [
                "Duracion del movimiento:    0.00 s",
                "Inicio movimiento:          n/d",
                "Fin movimiento:             n/d",
                "",
                "*** SIN MOVIMIENTO DETECTADO ***",
            ]
        head += [
            "",
            "Filtrado aplicado:          %s" % ("Si" if res.filter_info.get("applied") else "No"),
            "Metodo de filtrado:         %s" % res.filter_info.get("method", "-"),
            "Parametros del filtrado:    %s" % res.filter_info.get("params", "-"),
            "",
            "Estado:",
            "  %s" % ("Procesamiento completado correctamente"
                      if m["detected"] else
                      "Procesamiento completado - sin movimiento detectado"),
            "-" * 72,
            "",
            "",
        ]
        self.summary_widget.configure(state="normal")
        self.summary_widget.delete("1.0", "end")
        self.summary_widget.insert("1.0", "\n".join(head) + self.summary_text)
        self.summary_widget.see("1.0")
        self.summary_widget.configure(state="disabled")

    def _fill_tables(self):
        res = self.result
        self.window_var.set("Intervalo de metricas: %s" % res.metric_window_label())

        for name in AXIS_NAMES:
            e = res.joint_metrics[name]
            c = res.continuity_metrics[name]
            self.tree_main.insert("", "end", values=(
                name,
                _fmt(e["vel_max_abs_rad_s"], 5),
                _fmt(e["vel_rms_rad_s"], 5),
                _fmt(e["acc_max_abs_rad_s2"], 5),
                _fmt(e["acc_rms_rad_s2"], 5),
                _fmt(e["jerk_max_abs_rad_s3"], 5),
                _fmt(e["jerk_rms_rad_s3"], 5),
                _fmt(c["delta_q_max_rad"], 6),
            ))
            self.tree_pos.insert("", "end", values=(
                name,
                _fmt(e["pos_inicial_deg"], 4),
                _fmt(e["pos_final_deg"], 4),
                _fmt(e["pos_min_deg"], 4),
                _fmt(e["pos_max_deg"], 4),
                _fmt(e["recorrido_abs_deg"], 4),
            ))

        s = res.sampling
        m = res.motion
        rows = [
            ("Numero de muestras validas", "%d" % s["n_samples"], "-"),
            ("Filas del CSV original", "%d" % res.clean_report["rows_original"], "-"),
            ("Duracion total registrada", _fmt(s["duration_s"], 3), "s"),
            ("Duracion del movimiento",
             _fmt(m["duration_s"] if m["detected"] else 0.0, 3), "s"),
            ("Frecuencia efectiva", _fmt(s["freq_hz"], 3), "Hz"),
            ("Periodo medio", _fmt(s["dt_mean_ms"], 3), "ms"),
            ("Periodo mediano", _fmt(s["dt_median_ms"], 3), "ms"),
            ("Desviacion estandar del periodo", _fmt(s["dt_std_ms"], 3), "ms"),
            ("Periodo minimo", _fmt(s["dt_min_ms"], 3), "ms"),
            ("Periodo maximo", _fmt(s["dt_max_ms"], 3), "ms"),
            ("Movimiento detectado", "Si" if m["detected"] else "No", "-"),
            ("Umbral de velocidad efectivo",
             _fmt(m["effective_velocity_threshold"], 6), "rad/s"),
            ("Umbral de recorrido", _fmt(m["position_threshold"], 6), "rad"),
        ]
        for row in rows:
            self.tree_sampling.insert("", "end", values=row)

    def _embed_figures(self):
        for key, frame in self.plot_frames.items():
            fig = self.figures.get(key)
            if fig is None:
                reason = "Grafica no disponible."
                for fname, why in self.skipped_figures:
                    if FIG_FILENAMES.get(key) == fname:
                        reason = "Grafica no disponible: %s" % why
                ttk.Label(frame, text=reason, padding=20,
                          font=("TkDefaultFont", 11)).pack(anchor="center")
                continue
            fig.set_dpi(SCREEN_DPI)   # solo pantalla: el guardado usa PLOT_DPI
            canvas = FigureCanvasTkAgg(fig, master=frame)
            canvas.draw()
            toolbar = NavigationToolbar2Tk(canvas, frame, pack_toolbar=False)
            toolbar.update()
            toolbar.pack(side="bottom", fill="x")
            canvas.get_tk_widget().pack(side="top", fill="both", expand=True)
            self._canvases.append(canvas)


def main():
    if not os.path.isdir(RESULTS_DIR):
        try:
            os.makedirs(RESULTS_DIR, exist_ok=True)
        except OSError:
            pass
    try:
        app = ComparativeAnalyzerApp()
    except tk.TclError as exc:
        sys.stderr.write(
            "No se pudo iniciar la interfaz grafica (¿entorno sin pantalla?).\n"
            "Detalle: %s\n" % exc)
        return 1
    # Cierre limpio: cancela la reproduccion y apaga ROS2 si se inicializo.
    app.protocol("WM_DELETE_WINDOW", app.on_close)
    app.mainloop()
    return 0


if __name__ == "__main__":
    sys.exit(main())
